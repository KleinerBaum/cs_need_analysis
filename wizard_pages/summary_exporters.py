# wizard_pages/summary_exporters.py
"""Vertical Summary helpers extracted from ``08_summary.py``."""

import io
import hashlib
import json
from contextlib import nullcontext
from dataclasses import dataclass, replace
from collections import defaultdict
from html import escape
from typing import Any, Callable, Final, Mapping, Protocol, Sequence, TypedDict

import streamlit as st
import docx
from docx.image.image import Image as DocxImage

from constants import (
    INTAKE_FACTS,
    FactKey,
    FactRequirementStage,
    FactResolutionStatus,
    FactSourceType,
    NON_INTAKE_STEP_KEYS,
    SSKey,
    STEP_KEY_BENEFITS,
    STEP_KEY_COMPANY,
    STEP_KEY_INTERVIEW,
    STEP_KEY_ROLE_TASKS,
    STEP_KEY_SKILLS,
    SUMMARY_LOGO_UPLOAD_ALLOWED_EXTENSIONS,
    SUMMARY_LOGO_UPLOAD_ALLOWED_MIME_TYPES,
    SUMMARY_LOGO_UPLOAD_MAX_BYTES,
    SUMMARY_LOGO_UPLOAD_MIME_TYPE_BY_EXTENSION,
    UI_PREFERENCE_CONFIDENCE_THRESHOLD,
)
from interview_process import (
    build_candidate_stage_values,
    build_interview_export_payload,
    build_interview_value_rows,
    default_selected_interview_value_ids,
    normalize_interview_internal_flow,
)
from intake_facts import (
    build_intake_fact_resolution_state,
    get_intake_fact_evidence_state,
    get_intake_fact_state,
    latest_fact_confidence,
    mark_intake_facts_used_by_artifact,
    write_intake_fact,
)
from homepage_research import (
    normalize_company_website_research_payload as _normalize_company_website_research_payload,
)
from esco_client import EscoClient, EscoClientError
from esco_semantics import normalize_anchor_ref, sync_esco_semantic_state
from llm_client import (
    TASK_GENERATE_EMPLOYMENT_CONTRACT,
    JobAdGenerationResult,
    OpenAICallError,
    TASK_GENERATE_BOOLEAN_SEARCH,
    TASK_GENERATE_INTERVIEW_SHEET_HM,
    TASK_GENERATE_INTERVIEW_SHEET_HR,
    TASK_GENERATE_JOB_AD,
    TASK_GENERATE_VACANCY_BRIEF,
    generate_boolean_search_pack,
    generate_custom_job_ad,
    generate_employment_contract_draft,
    generate_interview_sheet_hm,
    generate_interview_sheet_hr,
    generate_vacancy_brief,
    resolve_model_for_task,
)
from schemas import (
    BooleanSearchPack,
    EscoConceptRef,
    EscoMatrixCoverageRow,
    EscoMappingReport,
    EscoSemanticContext,
    EscoUnresolvedTermDecision,
    EmploymentContractDraft,
    InterviewPrepSheetHiringManager,
    InterviewPrepSheetHR,
    JobAdExtract,
    LanguageRequirement,
    OccupationContextProfile,
    OccupationQuestionContext,
    QuestionFlowProvenance,
    QuestionPlan,
    QuestionStep,
    VacancyBrief,
    VacancyStructuredData,
    CompanyWebsiteResearch,
)
from settings_openai import load_openai_settings
from state import (
    clear_error,
    get_esco_occupation_selected,
    get_answers,
    get_answer_meta,
    has_confirmed_esco_anchor,
    get_model_override,
    handle_unexpected_exception,
    set_answer,
)
from question_dependencies import should_show_question
from question_limits import select_visible_questions_for_step_scope
from step_status import build_step_status_payload
from components.design_system import (
    render_card_start,
    render_critical_gaps,
    render_next_best_action,
    render_output_header,
    render_pill,
)
from summary_facts import (
    SummaryFactsRow,
    display_requirement_stage as _display_requirement_stage,
    display_salary_impact as _display_salary_impact,
    format_summary_answer_value as _format_summary_answer_value,
    format_summary_fact_value as _format_summary_fact_value,
    group_summary_fact_rows_by_area as _group_summary_fact_rows_by_area,
    status_for_answer_value as _status_for_answer_value,
    status_for_classification_value as _status_for_classification_value,
    status_for_value as _status_for_value,
    source_label_with_secondary_evidence as _source_label_with_secondary_evidence,
    summary_core_fact_row as _summary_core_fact_row,
    summary_fact_row_to_table_dict as _summary_fact_row_to_table_dict,
    summary_provenance_label as _summary_provenance_label,
)
from summary_artifacts import (
    artifact_display_label as _artifact_display_label,
    brief_pipeline_status_for_state,
    to_canonical_artifact_id as _to_canonical_artifact_id,
)
from summary_exports import (
    brief_to_markdown as _brief_to_markdown,
    build_summary_input_fingerprint as _build_summary_input_fingerprint,
)
from summary_esco import (
    build_esco_coverage_chart_spec as _build_esco_coverage_chart_spec,
    build_esco_coverage_kpis as _build_esco_coverage_kpis,
    build_esco_coverage_metrics as _build_esco_coverage_metrics,
    build_esco_mapping_report_csv as _build_esco_mapping_report_csv,
    extract_skills_step_raw_terms as _extract_skills_step_raw_terms,
    normalize_skill_term as _normalize_skill_term,
    to_esco_export_concepts as _to_esco_export_concepts,
)
from summary_job_ad import (
    dedupe_preserve_order as _dedupe_preserve_order,
    estimate_text_area_height as _estimate_text_area_height,
    job_ad_preview_html as _job_ad_preview_html_impl,
    job_ad_preview_shell_options as _job_ad_preview_shell_options_impl,
    job_ad_to_docx_bytes as _job_ad_to_docx_bytes_impl,
    job_ad_to_pdf_bytes as _job_ad_to_pdf_bytes_impl,
    sanitize_generated_job_ad as _sanitize_generated_job_ad,
)
from ui_components import (
    inject_pills_grid_css,
    render_boolean_risks,
    render_boolean_search_pack,
    render_boolean_supporting_terms,
    render_boolean_usage_notes,
    render_brief,
    render_employment_contract_draft,
    render_error_banner,
    render_interview_prep_fach,
    render_interview_prep_hr,
    render_openai_error,
)
from usage_events import get_usage_events, record_artifact_generated
from usage_utils import usage_has_cache_hit
from wizard_pages.base import (
    WizardContext,
    WizardPage,
    get_current_ui_mode,
    nav_buttons,
    render_active_ui_mode_caption,
)

SUPPORTED_LOGO_MIME_TYPES: dict[str, str] = {
    mime_type: "JPEG" if mime_type.endswith("jpeg") else "PNG"
    for mime_type in SUMMARY_LOGO_UPLOAD_ALLOWED_MIME_TYPES
}
SUPPORTED_LOGO_EXTENSIONS: tuple[str, ...] = SUMMARY_LOGO_UPLOAD_ALLOWED_EXTENSIONS


class DocxPictureDocument(Protocol):
    def add_picture(self, image_path_or_stream: Any, width: Any = ...) -> Any: ...


class LogoPayload(TypedDict):
    name: str
    mime_type: str
    extension: str
    byte_size: int
    width_px: int
    height_px: int
    bytes: bytes


class EscoExportConcept(TypedDict):
    uri: str
    label: str


class EscoMatchExplainability(TypedDict, total=False):
    reason: str
    confidence: str
    provenance_categories: list[str]


class EscoSharedFields(TypedDict):
    selected_occupation_uri: str
    essential_skills: list[dict[str, Any]]
    optional_skills: list[dict[str, Any]]
    unmapped_terms: list[str]
    unmapped_roles: list[str]
    unmapped_actions: dict[str, Any]


def _session_list(key: SSKey, default: list[Any] | None = None) -> list[Any]:
    raw = st.session_state.get(key.value, default if default is not None else [])
    return raw if isinstance(raw, list) else []


def _session_dict(key: SSKey, default: dict[str, Any] | None = None) -> dict[str, Any]:
    raw = st.session_state.get(key.value, default if default is not None else {})
    return raw if isinstance(raw, dict) else {}


def _session_str(key: SSKey, default: str = "") -> str:
    return str(st.session_state.get(key.value, default) or "").strip()


def _read_esco_shared_fields() -> EscoSharedFields:
    selected_uri = _session_str(SSKey.ESCO_SELECTED_OCCUPATION_URI)
    essential_raw = _session_list(
        SSKey.ESCO_CONFIRMED_ESSENTIAL_SKILLS,
        default=_session_list(SSKey.ESCO_SKILLS_SELECTED_MUST),
    )
    optional_raw = _session_list(
        SSKey.ESCO_CONFIRMED_OPTIONAL_SKILLS,
        default=_session_list(SSKey.ESCO_SKILLS_SELECTED_NICE),
    )
    unmapped_raw = _session_list(SSKey.ESCO_UNMAPPED_REQUIREMENT_TERMS)
    unmapped_role_raw = _session_list(SSKey.ESCO_UNMAPPED_ROLE_TERMS)
    unmapped_actions_raw = _session_dict(SSKey.ESCO_UNMAPPED_TERM_ACTIONS)
    essential = essential_raw if isinstance(essential_raw, list) else []
    optional = optional_raw if isinstance(optional_raw, list) else []
    unmapped = (
        [str(item).strip() for item in unmapped_raw if str(item).strip()]
        if isinstance(unmapped_raw, list)
        else []
    )
    unmapped_roles = (
        [str(item).strip() for item in unmapped_role_raw if str(item).strip()]
        if isinstance(unmapped_role_raw, list)
        else []
    )
    return {
        "selected_occupation_uri": selected_uri,
        "essential_skills": essential,
        "optional_skills": optional,
        "unmapped_terms": unmapped,
        "unmapped_roles": unmapped_roles,
        "unmapped_actions": unmapped_actions_raw,
    }


def _read_esco_semantic_context() -> EscoSemanticContext:
    selected_occupation = get_esco_occupation_selected()
    if isinstance(selected_occupation, dict):
        anchor = normalize_anchor_ref(selected_occupation)
        if anchor is not None:
            st.session_state[SSKey.ESCO_PRIMARY_ANCHOR.value] = anchor
    return sync_esco_semantic_state(st.session_state)


def _show_semantic_esco_sections() -> bool:
    try:
        semantic_context = _read_esco_semantic_context()
    except Exception:
        semantic_context = None
    return bool(
        semantic_context is not None
        and semantic_context.can_use_semantic_exports
    ) or has_confirmed_esco_anchor()


def _count_skill_relation_traces(skills: list[dict[str, Any]]) -> int:
    relation_traces = 0
    for item in skills:
        if not isinstance(item, dict):
            continue
        if str(item.get("related_occupation_uri") or "").strip():
            relation_traces += 1
            continue
        skill_uri = str(item.get("uri") or "").strip()
        if not skill_uri:
            continue
        try:
            payload = EscoClient().get_skill_related_occupations(
                skill_uri=skill_uri, limit=1
            )
        except EscoClientError:
            continue
        if isinstance(payload, dict) and payload:
            relation_traces += 1
    return relation_traces


def _compute_esco_coverage_metrics(shared_esco: EscoSharedFields) -> dict[str, int]:
    job_extract = st.session_state.get(SSKey.JOB_EXTRACT.value, {})
    return _build_esco_coverage_metrics(
        job_extract_payload=job_extract,
        essential_skills=shared_esco.get("essential_skills", []),
        optional_skills=shared_esco.get("optional_skills", []),
    )


def _build_esco_mapping_report_rows() -> list[dict[str, str]]:
    report_payload = st.session_state.get(SSKey.ESCO_SKILLS_MAPPING_REPORT.value, {})
    try:
        report = EscoMappingReport.model_validate(report_payload)
    except Exception:
        report = EscoMappingReport(
            mapped_count=0, unmapped_terms=[], collisions=[], notes=[]
        )

    shared_esco = _read_esco_shared_fields()
    chosen_must = _to_esco_export_concepts(shared_esco["essential_skills"])
    chosen_nice = _to_esco_export_concepts(shared_esco["optional_skills"])
    chosen_concepts = chosen_must + chosen_nice

    by_label: defaultdict[str, list[EscoExportConcept]] = defaultdict(list)
    for concept in chosen_concepts:
        by_label[_normalize_skill_term(concept.get("label", ""))].append(concept)

    raw_terms = _extract_skills_step_raw_terms(
        st.session_state.get(SSKey.JOB_EXTRACT.value, {})
    )

    rows: list[dict[str, str]] = []
    linked_uris: set[str] = set()
    normalized_unmapped = {
        _normalize_skill_term(term): term
        for term in (
            shared_esco["unmapped_terms"]
            if shared_esco["unmapped_terms"]
            else report.unmapped_terms
        )
    }
    for raw_term in raw_terms:
        label_matches = by_label.get(_normalize_skill_term(raw_term), [])
        if label_matches:
            for concept in label_matches:
                chosen_uri = concept.get("uri", "").strip()
                if chosen_uri:
                    linked_uris.add(chosen_uri)
                rows.append(
                    {
                        "raw_term": raw_term,
                        "chosen_uri": chosen_uri,
                        "chosen_label": concept.get("label", "").strip(),
                        "match_method": "label_exact",
                        "notes": "",
                    }
                )
        elif _normalize_skill_term(raw_term) in normalized_unmapped:
            rows.append(
                {
                    "raw_term": raw_term,
                    "chosen_uri": "",
                    "chosen_label": "",
                    "match_method": "unmapped",
                    "notes": "",
                }
            )

    for concept in chosen_concepts:
        chosen_uri = concept.get("uri", "").strip()
        if not chosen_uri or chosen_uri in linked_uris:
            continue
        rows.append(
            {
                "raw_term": "",
                "chosen_uri": chosen_uri,
                "chosen_label": concept.get("label", "").strip(),
                "match_method": "manual_selection",
                "notes": "",
            }
        )

    for term in report.collisions:
        rows.append(
            {
                "raw_term": str(term or "").strip(),
                "chosen_uri": "",
                "chosen_label": "",
                "match_method": "collision",
                "notes": "",
            }
        )
    for note in report.notes:
        rows.append(
            {
                "raw_term": "",
                "chosen_uri": "",
                "chosen_label": "",
                "match_method": "report_note",
                "notes": str(note or "").strip(),
            }
        )

    return sorted(
        rows,
        key=lambda row: (
            row["raw_term"].casefold(),
            row["chosen_uri"].casefold(),
            row["chosen_label"].casefold(),
            row["match_method"].casefold(),
            row["notes"].casefold(),
        ),
    )


def _build_structured_export_payload(brief: VacancyBrief) -> dict[str, Any]:
    payload = dict(brief.structured_data.model_dump(mode="json", exclude_none=True))
    intake_facts = get_intake_fact_state(st.session_state)
    if intake_facts:
        payload["intake_facts"] = dict(intake_facts)
        session_structured_fields = {
            "skill_items": FactKey.SKILLS_ITEMS,
            "variable_pay": FactKey.BENEFITS_VARIABLE_PAY,
            "travel_profile": FactKey.ROLE_TRAVEL_PROFILE,
            "interview_scorecard_template": FactKey.INTERVIEW_SCORECARD_TEMPLATE,
        }
        for payload_key, fact_key in session_structured_fields.items():
            if payload.get(payload_key):
                continue
            value = intake_facts.get(fact_key.value)
            if value:
                payload[payload_key] = value
    intake_fact_evidence = get_intake_fact_evidence_state(st.session_state)
    if intake_fact_evidence:
        payload["intake_fact_evidence"] = dict(intake_fact_evidence)
    intake_fact_resolution = build_intake_fact_resolution_state(st.session_state)
    if intake_fact_resolution:
        payload["intake_fact_resolution"] = intake_fact_resolution
    try:
        export_job = JobAdExtract.model_validate(brief.structured_data.job_extract)
    except Exception:
        export_job = JobAdExtract()
    export_answers = dict(brief.structured_data.answers)
    plan_payload = st.session_state.get(SSKey.QUESTION_PLAN.value)
    export_plan: QuestionPlan | None = None
    if isinstance(plan_payload, dict):
        try:
            export_plan = QuestionPlan.model_validate(plan_payload)
        except Exception:
            export_plan = None
    payload["interview_process"] = build_interview_export_payload(
        job=export_job,
        answers=export_answers,
        plan=export_plan,
        internal_flow=normalize_interview_internal_flow(
            st.session_state.get(SSKey.INTERVIEW_INTERNAL_FLOW.value, {})
        ),
    )
    occupation_profile_raw = st.session_state.get(SSKey.OCCUPATION_PROFILE.value)
    if isinstance(occupation_profile_raw, dict):
        try:
            payload["occupation_context_profile"] = (
                OccupationContextProfile.model_validate(
                    occupation_profile_raw
                ).model_dump(mode="json", exclude_none=True)
            )
        except Exception:
            pass
    question_context_raw = st.session_state.get(SSKey.OCCUPATION_QUESTION_CONTEXT.value)
    if isinstance(question_context_raw, dict):
        try:
            payload["occupation_question_context"] = (
                OccupationQuestionContext.model_validate(
                    question_context_raw
                ).model_dump(mode="json", exclude_none=True)
            )
        except Exception:
            pass
    flow_provenance_raw = st.session_state.get(SSKey.QUESTION_FLOW_PROVENANCE.value)
    if isinstance(flow_provenance_raw, dict) and flow_provenance_raw:
        try:
            payload["question_flow_provenance"] = (
                QuestionFlowProvenance.model_validate(
                    flow_provenance_raw
                ).model_dump(mode="json", exclude_none=True)
            )
        except Exception:
            pass
    semantic_context = _read_esco_semantic_context()
    payload["semantic_export_mode"] = semantic_context.semantic_export_mode
    payload["esco_anchor_state"] = semantic_context.anchor_state
    if semantic_context.capability_snapshot is not None:
        payload["esco_capability_snapshot"] = (
            semantic_context.capability_snapshot.model_dump(
                mode="json",
                exclude_none=True,
            )
        )

    selected_occupation = get_esco_occupation_selected()
    if semantic_context.can_use_semantic_exports:
        if semantic_context.primary_anchor is not None:
            primary_anchor = semantic_context.primary_anchor.model_dump(
                mode="json",
                exclude_none=True,
            )
            payload["esco_primary_anchor"] = primary_anchor
            payload["esco_occupations"] = [
                {
                    "uri": semantic_context.primary_anchor.uri,
                    "label": semantic_context.primary_anchor.title,
                }
            ]
            explainability = _read_esco_match_explainability()
            if explainability:
                payload["esco_occupation_provenance"] = explainability
        elif isinstance(selected_occupation, dict):
            try:
                parsed_occupation = EscoConceptRef.model_validate(selected_occupation)
            except Exception:
                pass
            else:
                payload["esco_occupations"] = [
                    {"uri": parsed_occupation.uri, "label": parsed_occupation.title}
                ]
                explainability = _read_esco_match_explainability()
                if explainability:
                    payload["esco_occupation_provenance"] = explainability
        if semantic_context.secondary_anchors:
            payload["esco_secondary_anchors"] = [
                anchor.model_dump(mode="json", exclude_none=True)
                for anchor in semantic_context.secondary_anchors
            ]

        shared_esco = _read_esco_shared_fields()
        must_skills = _to_esco_export_concepts(shared_esco["essential_skills"])
        if must_skills:
            payload["esco_skills_must"] = must_skills

        nice_skills = _to_esco_export_concepts(shared_esco["optional_skills"])
        if nice_skills:
            payload["esco_skills_nice"] = nice_skills
        if shared_esco["unmapped_terms"]:
            payload["esco_unmapped_requirement_terms"] = shared_esco["unmapped_terms"]
        if shared_esco["unmapped_roles"]:
            payload["esco_unmapped_role_terms"] = shared_esco["unmapped_roles"]
        unresolved_decisions_raw = _session_list(SSKey.ESCO_UNRESOLVED_TERM_DECISIONS)
        if shared_esco["unmapped_actions"]:
            payload["esco_unmapped_term_actions"] = shared_esco["unmapped_actions"]
        unresolved_decisions: list[dict[str, Any]] = []
        if isinstance(unresolved_decisions_raw, list):
            for entry in unresolved_decisions_raw:
                if not isinstance(entry, dict):
                    continue
                try:
                    parsed = EscoUnresolvedTermDecision.model_validate(entry)
                except ValueError:
                    continue
                unresolved_decisions.append(
                    parsed.model_dump(mode="json", exclude_none=True)
                )
        if not unresolved_decisions and shared_esco["unmapped_actions"]:
            for raw_term, action_payload in shared_esco["unmapped_actions"].items():
                if not isinstance(action_payload, dict):
                    continue
                candidate = dict(action_payload)
                candidate.setdefault("raw_term", str(raw_term))
                try:
                    parsed = EscoUnresolvedTermDecision.model_validate(candidate)
                except ValueError:
                    continue
                unresolved_decisions.append(
                    parsed.model_dump(mode="json", exclude_none=True)
                )
        if unresolved_decisions:
            payload["esco_unresolved_term_decisions"] = unresolved_decisions

    esco_config = _session_dict(SSKey.ESCO_CONFIG)
    release_lane = str(
        esco_config.get("release_lane")
        or st.session_state.get(SSKey.ESCO_RELEASE_LANE.value)
        or ""
    ).strip()
    if release_lane:
        payload["esco_release_lane"] = release_lane
    selected_version = str(esco_config.get("selected_version") or "").strip()
    if selected_version:
        payload["esco_version"] = selected_version
    esco_source = str(st.session_state.get(SSKey.ESCO_LAST_DATA_SOURCE.value) or "").strip()
    if esco_source:
        payload["esco_data_source"] = esco_source
    data_source_mode = str(esco_config.get("data_source_mode") or "").strip()
    if data_source_mode:
        payload["esco_data_source_mode"] = data_source_mode
    if semantic_context.can_use_matrix_coverage:
        matrix_metadata = st.session_state.get(SSKey.ESCO_MATRIX_METADATA.value, {})
        matrix_loaded = bool(st.session_state.get(SSKey.ESCO_MATRIX_LOADED.value, False))
        if isinstance(matrix_metadata, dict) and (matrix_loaded or matrix_metadata):
            payload["esco_matrix"] = {
                "enabled": bool(
                    st.session_state.get(SSKey.ESCO_MATRIX_ENABLED.value, False)
                ),
                "loaded": matrix_loaded,
                "source": str(matrix_metadata.get("source") or "").strip(),
                "version": str(matrix_metadata.get("version") or "").strip(),
                "records": int(matrix_metadata.get("records") or 0),
            }
        matrix_coverage_rows_raw = _session_list(SSKey.ESCO_MATRIX_COVERAGE_ROWS)
        matrix_coverage_rows: list[dict[str, Any]] = []
        for row in matrix_coverage_rows_raw:
            if not isinstance(row, dict):
                continue
            try:
                matrix_coverage_rows.append(
                    EscoMatrixCoverageRow.model_validate(row).model_dump(mode="json")
                )
            except Exception:
                continue
        if matrix_coverage_rows:
            payload["esco_matrix_coverage"] = matrix_coverage_rows
            if isinstance(payload.get("esco_matrix"), dict):
                payload["esco_matrix"]["coverage_rows"] = len(matrix_coverage_rows)
        matrix_coverage_context = _session_dict(SSKey.ESCO_MATRIX_COVERAGE_CONTEXT)
        if matrix_coverage_context:
            context_payload = {
                "reason": str(matrix_coverage_context.get("reason") or "").strip(),
                "occupation_group": str(
                    matrix_coverage_context.get("occupation_group") or ""
                ).strip(),
                "rows": int(matrix_coverage_context.get("rows") or 0),
            }
            if any(context_payload.values()):
                payload["esco_matrix_coverage_context"] = context_payload

        title_variants_raw = _session_dict(SSKey.ESCO_OCCUPATION_TITLE_VARIANTS)
        variants_uri = str(title_variants_raw.get("uri") or "").strip()
        recommended_titles_raw = title_variants_raw.get("recommended_titles", {})
        if (
            isinstance(selected_occupation, dict)
            and variants_uri == str(selected_occupation.get("uri") or "").strip()
            and isinstance(recommended_titles_raw, dict)
        ):
            recommended_titles: dict[str, list[str]] = {}
            for language, labels_raw in recommended_titles_raw.items():
                if not isinstance(language, str) or not isinstance(labels_raw, list):
                    continue
                labels = [
                    str(label).strip()
                    for label in labels_raw
                    if isinstance(label, str) and str(label).strip()
                ]
                if labels:
                    recommended_titles[language] = labels
            if recommended_titles:
                payload["recommended_titles"] = recommended_titles

    selected_benefits = _read_saved_selection_labels(SSKey.BENEFITS_SELECTED)
    if selected_benefits:
        payload["selected_benefits"] = selected_benefits

    salary_forecast = st.session_state.get(SSKey.SALARY_FORECAST_LAST_RESULT.value)
    if isinstance(salary_forecast, dict):
        payload["salary_forecast"] = salary_forecast

    scenario_lab_rows = st.session_state.get(SSKey.SALARY_SCENARIO_LAB_ROWS.value)
    if isinstance(scenario_lab_rows, list):
        payload["salary_scenarios"] = scenario_lab_rows[:100]

    def _normalized_provenance_rows(raw_rows: Any) -> list[dict[str, Any]]:
        if not isinstance(raw_rows, list):
            return []
        rows: list[dict[str, Any]] = []
        for item in raw_rows:
            if not isinstance(item, dict):
                continue
            label = str(item.get("label") or "").strip()
            if not label:
                continue
            row = {
                "label": label,
                "source_hint": str(item.get("source_hint") or "").strip(),
                "source_file": str(item.get("source_file") or "").strip(),
                "concept_uri": str(item.get("concept_uri") or item.get("uri") or "").strip(),
                "rationale": str(item.get("rationale") or "").strip(),
                "evidence": str(item.get("evidence") or "").strip(),
            }
            if not row["source_hint"]:
                source_text = str(item.get("source") or "").strip().casefold()
                row["source_hint"] = "esco_rag" if "rag" in source_text else "llm"
            rows.append(row)
        return rows

    role_task_suggestions = _normalized_provenance_rows(
        st.session_state.get(SSKey.ROLE_TASKS_LLM_SUGGESTED.value, [])
    )
    if role_task_suggestions:
        payload["role_task_suggestions"] = role_task_suggestions

    skill_suggestions = _normalized_provenance_rows(
        st.session_state.get(SSKey.SKILLS_LLM_SUGGESTED.value, [])
    )
    if skill_suggestions:
        payload["skill_suggestions"] = skill_suggestions

    benefit_suggestions = _normalized_provenance_rows(
        st.session_state.get(SSKey.BENEFITS_LLM_SUGGESTED.value, [])
    )
    if benefit_suggestions:
        payload["benefit_suggestions"] = benefit_suggestions
    return payload


def _build_brief_structured_preview_payload(brief: VacancyBrief) -> dict[str, Any]:
    export_payload = _build_structured_export_payload(brief)
    preview_keys = (
        "job_extract",
        "question_plan",
        "answers",
        "selected_role_tasks",
        "selected_skills",
        "selected_benefits",
    )
    return {
        key: export_payload[key]
        for key in preview_keys
        if key in export_payload
    }


def _read_saved_selection_labels(key: SSKey) -> list[str]:
    raw = st.session_state.get(key.value, [])
    if not isinstance(raw, list):
        return []
    values: list[str] = []
    for item in raw:
        if not isinstance(item, str):
            continue
        label = item.strip()
        if label:
            values.append(label)
    return values


def _read_esco_skill_refs(key: SSKey) -> list[dict[str, str]]:
    raw_value = st.session_state.get(key.value, [])
    if not isinstance(raw_value, list):
        return []
    refs: list[dict[str, str]] = []
    for item in raw_value:
        if not isinstance(item, dict):
            continue
        uri = str(item.get("uri", "") or "").strip()
        title = str(item.get("title", "") or "").strip()
        if not uri and not title:
            continue
        refs.append({"uri": uri, "title": title})
    return refs


def _read_selected_esco_occupation() -> dict[str, str]:
    shared_esco = _read_esco_shared_fields()
    selected_occupation = get_esco_occupation_selected()
    if not isinstance(selected_occupation, dict):
        if shared_esco["selected_occupation_uri"]:
            return {"uri": shared_esco["selected_occupation_uri"], "title": ""}
        return {}
    return {
        "uri": shared_esco["selected_occupation_uri"]
        or str(selected_occupation.get("uri", "") or "").strip(),
        "title": str(selected_occupation.get("title", "") or "").strip(),
    }


def _read_esco_match_explainability() -> EscoMatchExplainability:
    reason_raw = st.session_state.get(SSKey.ESCO_MATCH_REASON.value)
    confidence_raw = st.session_state.get(SSKey.ESCO_MATCH_CONFIDENCE.value)
    provenance_raw = _session_list(SSKey.ESCO_MATCH_PROVENANCE)

    reason = str(reason_raw or "").strip()
    confidence = str(confidence_raw or "").strip()
    provenance = [str(item).strip() for item in provenance_raw if str(item).strip()]

    explainability: EscoMatchExplainability = {}
    if reason:
        explainability["reason"] = reason
    if confidence:
        explainability["confidence"] = confidence
    if provenance:
        explainability["provenance_categories"] = provenance
    return explainability


def _job_ad_to_docx_bytes(
    job_ad: JobAdGenerationResult,
    styleguide: str = "",
    *,
    logo_payload: LogoPayload | None = None,
) -> bytes:
    if logo_payload is None:
        logo_payload = _read_logo_payload()
    return _job_ad_to_docx_bytes_impl(
        job_ad,
        styleguide=styleguide,
        logo_payload=logo_payload,
    )


def _job_ad_to_pdf_bytes(
    job_ad: JobAdGenerationResult,
    styleguide: str = "",
    *,
    logo_payload: LogoPayload | None = None,
) -> bytes | None:
    if logo_payload is None:
        logo_payload = _read_logo_payload()
    return _job_ad_to_pdf_bytes_impl(
        job_ad,
        styleguide=styleguide,
        logo_payload=logo_payload,
    )


def _brief_to_docx_bytes(brief: VacancyBrief) -> bytes:
    d = docx.Document()
    logo_payload = _read_logo_payload()
    _add_logo_to_docx(document=d, logo_payload=logo_payload)
    d.add_heading("Recruiting Brief", level=1)
    d.add_paragraph(f"One-liner: {brief.one_liner}")

    d.add_heading("Hiring Context", level=2)
    d.add_paragraph(brief.hiring_context)

    d.add_heading("Role Summary", level=2)
    d.add_paragraph(brief.role_summary)

    d.add_heading("Top Responsibilities", level=2)
    for x in brief.top_responsibilities:
        d.add_paragraph(x, style="List Bullet")

    d.add_heading("Must-have", level=2)
    for x in brief.must_have:
        d.add_paragraph(x, style="List Bullet")

    d.add_heading("Nice-to-have", level=2)
    for x in brief.nice_to_have:
        d.add_paragraph(x, style="List Bullet")

    d.add_heading("Interview Plan", level=2)
    for x in brief.interview_plan:
        d.add_paragraph(x, style="List Bullet")

    d.add_heading("Risks / Open Questions", level=2)
    for x in brief.risks_open_questions:
        d.add_paragraph(x, style="List Bullet")

    d.add_heading("Job Ad Draft (DE)", level=2)
    d.add_paragraph(brief.job_ad_draft)

    bio = io.BytesIO()
    d.save(bio)
    return bio.getvalue()


def _interview_prep_hr_to_docx_bytes(sheet: InterviewPrepSheetHR) -> bytes:
    d = docx.Document()
    d.add_heading("Interview Sheet (HR)", level=1)
    d.add_paragraph(f"Rolle: {sheet.role_title}")
    d.add_paragraph(f"Interview-Stage: {sheet.interview_stage}")
    d.add_paragraph(f"Dauer: {sheet.duration_minutes} Minuten")

    d.add_heading("Opening Script", level=2)
    d.add_paragraph(sheet.opening_script)

    d.add_heading("Frageblöcke", level=2)
    for block in sheet.question_blocks:
        d.add_heading(block.title, level=3)
        d.add_paragraph(f"Ziel: {block.objective}")
        if block.questions:
            d.add_paragraph("Fragen:")
            for question in block.questions:
                d.add_paragraph(question, style="List Bullet")
        if block.follow_up_prompts:
            d.add_paragraph("Follow-ups:")
            for prompt in block.follow_up_prompts:
                d.add_paragraph(prompt, style="List Bullet")

    d.add_heading("Knockout-Kriterien", level=2)
    for knockout_criterion in sheet.knockout_criteria:
        d.add_paragraph(knockout_criterion, style="List Bullet")

    d.add_heading("Bewertungsrubrik", level=2)
    for rubric_criterion in sheet.evaluation_rubric:
        d.add_heading(rubric_criterion.title, level=3)
        d.add_paragraph(rubric_criterion.description)
        d.add_paragraph(f"Gewichtung: {rubric_criterion.weight_percent} %")
        if rubric_criterion.score_scale:
            d.add_paragraph(f"Skala: {' | '.join(rubric_criterion.score_scale)}")
        if rubric_criterion.evidence_examples:
            d.add_paragraph("Evidenz-Beispiele:")
            for evidence in rubric_criterion.evidence_examples:
                d.add_paragraph(evidence, style="List Bullet")

    d.add_heading("Finale Empfehlung", level=2)
    for option in sheet.final_recommendation_options:
        d.add_paragraph(option, style="List Bullet")

    bio = io.BytesIO()
    d.save(bio)
    return bio.getvalue()


def _apply_interview_docx_brand_style(document: docx.Document, styleguide: str) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = docx.shared.Pt(10.5)
    for style_name in ("Title", "Heading 1", "Heading 2", "Heading 3"):
        style = document.styles[style_name]
        style.font.name = "Arial"
    if styleguide.strip():
        for style_name in ("Heading 1", "Heading 2"):
            document.styles[style_name].font.color.rgb = docx.shared.RGBColor(
                31, 78, 121
            )


def _interview_prep_fach_to_docx_bytes(
    sheet: InterviewPrepSheetHiringManager,
    *,
    logo_payload: LogoPayload | None = None,
    styleguide: str = "",
) -> bytes:
    d = docx.Document()
    _apply_interview_docx_brand_style(d, styleguide)
    if logo_payload is None:
        logo_payload = _read_logo_payload()
    _add_logo_to_docx(document=d, logo_payload=logo_payload)
    d.add_heading("Interview Sheet (Fachbereich)", level=1)
    d.add_paragraph(f"Rolle: {sheet.role_title}")
    d.add_paragraph(f"Interview-Stage: {sheet.interview_stage}")
    d.add_paragraph(f"Dauer: {sheet.duration_minutes} Minuten")

    d.add_heading("Kompetenzen validieren", level=2)
    for competency in sheet.competencies_to_validate:
        d.add_paragraph(competency, style="List Bullet")

    d.add_heading("Frageblöcke", level=2)
    for block in sheet.question_blocks:
        d.add_heading(block.title, level=3)
        d.add_paragraph(f"Ziel: {block.objective}")
        if block.questions:
            d.add_paragraph("Fragen:")
            for question in block.questions:
                d.add_paragraph(question, style="List Bullet")
        if block.follow_up_prompts:
            d.add_paragraph("Follow-ups:")
            for follow_up in block.follow_up_prompts:
                d.add_paragraph(follow_up, style="List Bullet")

    d.add_heading("Technical Deep Dive", level=2)
    for topic in sheet.technical_deep_dive_topics:
        d.add_paragraph(topic, style="List Bullet")

    d.add_heading("Case/Task Prompt", level=2)
    d.add_paragraph(sheet.case_or_task_prompt or "Kein Case/Task hinterlegt.")

    d.add_heading("Bewertungsrubrik", level=2)
    for criterion in sheet.evaluation_rubric:
        d.add_heading(criterion.title, level=3)
        d.add_paragraph(criterion.description)
        d.add_paragraph(f"Gewichtung: {criterion.weight_percent} %")
        if criterion.score_scale:
            d.add_paragraph(f"Skala: {' | '.join(criterion.score_scale)}")
        if criterion.evidence_examples:
            d.add_paragraph("Evidenz-Beispiele:")
            for evidence in criterion.evidence_examples:
                d.add_paragraph(evidence, style="List Bullet")

    d.add_heading("Debrief-Fragen", level=2)
    for question in sheet.debrief_questions:
        d.add_paragraph(question, style="List Bullet")

    bio = io.BytesIO()
    d.save(bio)
    return bio.getvalue()


def _interview_prep_fach_to_pdf_bytes(
    sheet: InterviewPrepSheetHiringManager,
    *,
    logo_payload: LogoPayload | None = None,
    styleguide: str = "",
) -> bytes | None:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import cm
        from reportlab.lib.utils import ImageReader
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib.colors import HexColor
        from reportlab.platypus import (
            Image,
            ListFlowable,
            ListItem,
            Paragraph,
            SimpleDocTemplate,
            Spacer,
        )
    except Exception:
        return None

    bio = io.BytesIO()
    document = SimpleDocTemplate(
        bio,
        pagesize=A4,
        leftMargin=2 * cm,
        rightMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
        title="Interview Sheet (Fachbereich)",
        author="anonymous",
    )
    styles = getSampleStyleSheet()
    if styleguide.strip():
        styles["Heading1"].textColor = HexColor("#1F4E79")
        styles["Heading2"].textColor = HexColor("#1F4E79")
    story: list[Any] = []

    if logo_payload is None:
        logo_payload = _read_logo_payload()
    if logo_payload is not None:
        logo_bytes = logo_payload.get("bytes")
        if isinstance(logo_bytes, bytes) and logo_bytes:
            try:
                image_width, image_height = ImageReader(io.BytesIO(logo_bytes)).getSize()
                max_width = 4.2 * cm
                max_height = 1.8 * cm
                scale = min(max_width / image_width, max_height / image_height, 1)
                story.append(
                    Image(
                        io.BytesIO(logo_bytes),
                        width=image_width * scale,
                        height=image_height * scale,
                    )
                )
                story.append(Spacer(1, 0.5 * cm))
            except Exception:
                pass

    def _paragraph(value: str, style_name: str = "BodyText") -> Paragraph:
        return Paragraph(escape(value).replace("\n", "<br/>"), styles[style_name])

    def _heading(value: str, style_name: str = "Heading2") -> None:
        if value.strip():
            story.append(_paragraph(value.strip(), style_name))

    def _bullets(items: Sequence[str]) -> None:
        clean_items = _dedupe_preserve_order(
            [str(item).strip() for item in items if str(item).strip()]
        )
        if not clean_items:
            return
        story.append(
            ListFlowable(
                [ListItem(_paragraph(item), leftIndent=0) for item in clean_items],
                bulletType="bullet",
                leftIndent=14,
            )
        )

    _heading("Interview Sheet (Fachbereich)", "Title")
    story.append(_paragraph(f"Rolle: {sheet.role_title}"))
    story.append(_paragraph(f"Interview-Stage: {sheet.interview_stage}"))
    story.append(_paragraph(f"Dauer: {sheet.duration_minutes} Minuten"))
    story.append(Spacer(1, 0.25 * cm))

    _heading("Kompetenzen validieren")
    _bullets(sheet.competencies_to_validate)
    _heading("Frageblöcke")
    for block in sheet.question_blocks:
        _heading(block.title, "Heading3")
        story.append(_paragraph(f"Ziel: {block.objective}"))
        if block.questions:
            story.append(_paragraph("Fragen:"))
            _bullets(block.questions)
        if block.follow_up_prompts:
            story.append(_paragraph("Follow-ups:"))
            _bullets(block.follow_up_prompts)
    _heading("Technical Deep Dive")
    _bullets(sheet.technical_deep_dive_topics)
    _heading("Case/Task Prompt")
    story.append(_paragraph(sheet.case_or_task_prompt or "Kein Case/Task hinterlegt."))
    _heading("Bewertungsrubrik")
    for criterion in sheet.evaluation_rubric:
        _heading(criterion.title, "Heading3")
        story.append(_paragraph(criterion.description))
        story.append(_paragraph(f"Gewichtung: {criterion.weight_percent} %"))
        if criterion.score_scale:
            story.append(_paragraph(f"Skala: {' | '.join(criterion.score_scale)}"))
        _bullets(criterion.evidence_examples)
    _heading("Debrief-Fragen")
    _bullets(sheet.debrief_questions)

    document.build(story)
    return bio.getvalue()


def _employment_contract_to_docx_bytes(draft: EmploymentContractDraft) -> bytes:
    d = docx.Document()
    d.add_heading("Arbeitsvertrag (Template Draft)", level=1)
    d.add_paragraph(
        "Hinweis: Diese Vorlage ist kein finaler Vertrag und keine Rechtsberatung."
    )
    d.add_paragraph(f"Jurisdiction: {draft.jurisdiction}")
    d.add_paragraph(f"Rolle: {draft.role_title}")
    d.add_paragraph(f"Employment Type: {draft.employment_type}")
    d.add_paragraph(f"Contract Type: {draft.contract_type}")
    d.add_paragraph(f"Start Date: {draft.start_date or '—'}")
    d.add_paragraph(
        "Probation (Monate): "
        + (str(draft.probation_period_months) if draft.probation_period_months else "—")
    )
    salary = draft.salary
    d.add_paragraph(
        "Salary: "
        f"{salary.min if salary.min is not None else '—'} - "
        f"{salary.max if salary.max is not None else '—'} "
        f"{salary.currency or ''} / {salary.period or ''}".strip()
    )
    if salary.notes:
        d.add_paragraph(f"Salary Notes: {salary.notes}")
    d.add_paragraph(
        "Working Hours / Week: "
        + (str(draft.working_hours_per_week) if draft.working_hours_per_week else "—")
    )
    d.add_paragraph(
        "Vacation Days / Year: "
        + (str(draft.vacation_days_per_year) if draft.vacation_days_per_year else "—")
    )
    d.add_paragraph(f"Place of Work: {draft.place_of_work or '—'}")
    d.add_paragraph(f"Notice Period: {draft.notice_period or '—'}")

    d.add_heading("Missing Inputs", level=2)
    if draft.missing_inputs:
        for missing_input in draft.missing_inputs:
            d.add_paragraph(missing_input, style="List Bullet")
    else:
        d.add_paragraph("Keine fehlenden Inputs markiert.")

    d.add_heading("Clauses", level=2)
    if draft.clauses:
        for clause in draft.clauses:
            d.add_heading(clause.title, level=3)
            d.add_paragraph(clause.clause_text)
            d.add_paragraph(f"Pflichtklausel: {'Ja' if clause.required else 'Nein'}")
            if clause.legal_note:
                d.add_paragraph(f"Legal Note: {clause.legal_note}")
    else:
        d.add_paragraph("Keine Klauseln hinterlegt.")

    d.add_heading("Signature Requirements", level=2)
    for requirement in draft.signature_requirements:
        d.add_paragraph(requirement, style="List Bullet")

    bio = io.BytesIO()
    d.save(bio)
    return bio.getvalue()


def _logo_basename(name: Any) -> str:
    raw_name = str(name or "").replace("\x00", "").replace("\\", "/").strip()
    return raw_name.rsplit("/", 1)[-1].strip()


def _logo_extension_from_name(name: Any) -> str:
    lower_name = _logo_basename(name).lower()
    for extension in sorted(
        SUMMARY_LOGO_UPLOAD_ALLOWED_EXTENSIONS,
        key=len,
        reverse=True,
    ):
        if lower_name.endswith(extension):
            return extension
    return ""


def _safe_logo_filename(name: Any, extension: str) -> str:
    basename = _logo_basename(name)
    stem = (
        basename[: -len(extension)]
        if basename.lower().endswith(extension)
        else basename
    )
    safe_stem = "".join(
        char if char.isprintable() and char not in {"/", "\\"} else "_"
        for char in stem
    ).strip(" ._")
    return f"{(safe_stem or 'logo')[:120]}{extension}"


def _coerce_logo_bytes(raw_bytes: Any) -> bytes | None:
    if isinstance(raw_bytes, memoryview):
        raw_bytes = raw_bytes.tobytes()
    if not isinstance(raw_bytes, (bytes, bytearray)):
        return None
    logo_bytes = bytes(raw_bytes)
    if not logo_bytes or len(logo_bytes) > SUMMARY_LOGO_UPLOAD_MAX_BYTES:
        return None
    return logo_bytes


def _decode_logo_image(logo_bytes: bytes) -> tuple[str, int, int] | None:
    try:
        image = DocxImage.from_blob(logo_bytes)
        mime_type = str(image.content_type or "").strip().lower()
        width_px = int(image.px_width or 0)
        height_px = int(image.px_height or 0)
    except Exception:
        return None
    if (
        mime_type not in SUMMARY_LOGO_UPLOAD_ALLOWED_MIME_TYPES
        or width_px <= 0
        or height_px <= 0
    ):
        return None
    return mime_type, width_px, height_px


def _normalize_logo_payload_data(
    *,
    name: Any,
    mime_type: Any,
    raw_bytes: Any,
) -> LogoPayload | None:
    normalized_mime_type = str(mime_type or "").lower().strip()
    if normalized_mime_type not in SUMMARY_LOGO_UPLOAD_ALLOWED_MIME_TYPES:
        return None
    extension = _logo_extension_from_name(name)
    if (
        not extension
        or SUMMARY_LOGO_UPLOAD_MIME_TYPE_BY_EXTENSION.get(extension)
        != normalized_mime_type
    ):
        return None
    logo_bytes = _coerce_logo_bytes(raw_bytes)
    if logo_bytes is None:
        return None
    decoded_logo = _decode_logo_image(logo_bytes)
    if decoded_logo is None:
        return None
    decoded_mime_type, width_px, height_px = decoded_logo
    if decoded_mime_type != normalized_mime_type:
        return None
    return {
        "name": _safe_logo_filename(name, extension),
        "mime_type": normalized_mime_type,
        "extension": extension,
        "byte_size": len(logo_bytes),
        "width_px": width_px,
        "height_px": height_px,
        "bytes": logo_bytes,
    }


def _normalize_logo_payload(uploaded_logo: Any) -> LogoPayload | None:
    if uploaded_logo is None:
        return None
    try:
        raw_bytes = uploaded_logo.getvalue()
    except Exception:
        return None
    return _normalize_logo_payload_data(
        name=getattr(uploaded_logo, "name", ""),
        mime_type=getattr(uploaded_logo, "type", ""),
        raw_bytes=raw_bytes,
    )


def _read_logo_payload() -> LogoPayload | None:
    raw_payload = st.session_state.get(SSKey.SUMMARY_LOGO.value)
    return _normalize_stored_logo_payload(raw_payload)


def _normalize_stored_logo_payload(raw_payload: Any) -> LogoPayload | None:
    if isinstance(raw_payload, dict):
        name = raw_payload.get("name") or ""
        if (
            not name
            and raw_payload.get("extension") in SUMMARY_LOGO_UPLOAD_ALLOWED_EXTENSIONS
        ):
            name = f"logo{raw_payload.get('extension')}"
        return _normalize_logo_payload_data(
            name=name,
            mime_type=raw_payload.get("mime_type"),
            raw_bytes=raw_payload.get("bytes"),
        )
    return _normalize_logo_payload(raw_payload)


def _job_ad_logo_payload(custom_job_ad_raw: Mapping[str, Any]) -> LogoPayload | None:
    if "logo" in custom_job_ad_raw:
        return _normalize_stored_logo_payload(custom_job_ad_raw.get("logo"))
    return _read_logo_payload()


def _add_logo_to_docx(
    document: DocxPictureDocument, logo_payload: LogoPayload | None
) -> bool:
    if logo_payload is None:
        return False
    logo_bytes = logo_payload.get("bytes")
    if not isinstance(logo_bytes, bytes):
        return False
    image_stream = io.BytesIO(bytes(logo_bytes))
    image_stream.seek(0)
    try:
        document.add_picture(image_stream, width=docx.shared.Cm(4.0))
    except Exception:
        return False
    return True


def _job_ad_preview_shell_options(options: Mapping[str, Any] | None) -> dict[str, Any]:
    return _job_ad_preview_shell_options_impl(options)


def _job_ad_preview_html(
    job_ad: JobAdGenerationResult,
    *,
    logo_payload: LogoPayload | None,
) -> str:
    return _job_ad_preview_html_impl(job_ad, logo_payload=logo_payload)
