"""Pure helpers for generated Summary job ads."""

from __future__ import annotations

import io
import re
from html import escape
from typing import Any, Mapping, Sequence

import docx

from document_preview import markdown_article_preview_html
from llm_client import JobAdGenerationResult


MARKDOWN_TOKEN_RE = re.compile(r"(\*\*|__|`|^#{1,6}\s*)", re.MULTILINE)
LogoPayload = Mapping[str, Any]


def normalize_list_item(value: str) -> str:
    return re.sub(r"^[\-•*\d\.)\s]+", "", value).strip()


def strip_inline_markdown(value: str) -> str:
    normalized = MARKDOWN_TOKEN_RE.sub("", value)
    normalized = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", normalized)
    return normalized.strip()


def dedupe_preserve_order(values: list[str]) -> list[str]:
    seen: set[str] = set()
    result: list[str] = []
    for item in values:
        normalized = strip_inline_markdown(item).strip()
        if not normalized:
            continue
        key = normalized.casefold()
        if key in seen:
            continue
        seen.add(key)
        result.append(normalized)
    return result


def sanitize_generated_job_ad(
    job_ad: JobAdGenerationResult,
) -> tuple[JobAdGenerationResult, list[str]]:
    body_lines: list[str] = []
    extracted_target_group: list[str] = []
    extracted_checklist: list[str] = []
    extracted_notes: list[str] = []

    section = "body"
    for raw_line in job_ad.job_ad_text.splitlines():
        line = raw_line.strip()
        if not line:
            if section == "body":
                body_lines.append("")
            continue

        lowered = line.rstrip(":").strip().casefold()
        if lowered == "zielgruppe":
            section = "target_group"
            continue
        if lowered in {"agg-checkliste", "agg checkliste"}:
            section = "agg_checklist"
            continue

        if line.casefold().startswith("hinweis:"):
            extracted_notes.append(line.split(":", 1)[1].strip())
            continue

        normalized_item = normalize_list_item(line)
        if section == "target_group":
            if normalized_item:
                extracted_target_group.append(normalized_item)
            continue
        if section == "agg_checklist":
            if normalized_item:
                extracted_checklist.append(normalized_item)
            continue

        body_lines.append(raw_line.rstrip())

    while body_lines and not body_lines[-1].strip():
        body_lines.pop()

    normalized_job_ad = JobAdGenerationResult(
        headline=strip_inline_markdown(job_ad.headline),
        target_group=dedupe_preserve_order(
            [*job_ad.target_group, *extracted_target_group]
        ),
        agg_checklist=dedupe_preserve_order(
            [*job_ad.agg_checklist, *extracted_checklist]
        ),
        job_ad_text=strip_inline_markdown("\n".join(body_lines).strip()),
        intro=strip_inline_markdown(job_ad.intro),
        responsibilities=dedupe_preserve_order(job_ad.responsibilities),
        profile=dedupe_preserve_order(job_ad.profile),
        offer=dedupe_preserve_order(job_ad.offer),
        cta=strip_inline_markdown(job_ad.cta),
        equal_opportunity_note=strip_inline_markdown(job_ad.equal_opportunity_note),
    )
    return normalized_job_ad, dedupe_preserve_order(extracted_notes)


def has_structured_job_ad_sections(job_ad: JobAdGenerationResult) -> bool:
    return any(
        (
            job_ad.intro.strip(),
            job_ad.responsibilities,
            job_ad.profile,
            job_ad.offer,
            job_ad.cta.strip(),
            job_ad.equal_opportunity_note.strip(),
        )
    )


def build_publishable_job_ad_markdown(job_ad: JobAdGenerationResult) -> str:
    if not has_structured_job_ad_sections(job_ad):
        return strip_inline_markdown(job_ad.job_ad_text)

    lines: list[str] = []
    headline = strip_inline_markdown(job_ad.headline)
    if headline:
        lines.extend([f"# {headline}", ""])
    if job_ad.intro.strip():
        lines.extend([strip_inline_markdown(job_ad.intro), ""])

    for heading, items in (
        ("Deine Aufgaben", job_ad.responsibilities),
        ("Dein Profil", job_ad.profile),
        ("Was wir bieten", job_ad.offer),
    ):
        clean_items = dedupe_preserve_order(items)
        if not clean_items:
            continue
        lines.extend([f"## {heading}", ""])
        lines.extend(f"- {item}" for item in clean_items)
        lines.append("")

    for value in (job_ad.cta, job_ad.equal_opportunity_note):
        clean_value = strip_inline_markdown(value)
        if clean_value:
            lines.extend([clean_value, ""])

    while lines and not lines[-1]:
        lines.pop()
    return "\n".join(lines).strip()


def build_publishable_job_ad_plain_text(job_ad: JobAdGenerationResult) -> str:
    markdown = build_publishable_job_ad_markdown(job_ad)
    lines: list[str] = []
    for line in markdown.splitlines():
        if line.startswith("# "):
            lines.append(line[2:].strip())
        elif line.startswith("## "):
            lines.append(line[3:].strip())
        else:
            lines.append(line)
    return "\n".join(lines).strip()


def _add_logo_to_docx(document: Any, logo_payload: LogoPayload | None) -> bool:
    if logo_payload is None:
        return False
    logo_bytes = logo_payload.get("bytes")
    if not isinstance(logo_bytes, bytes):
        return False
    image_stream = io.BytesIO(bytes(logo_bytes))
    image_stream.seek(0)
    try:
        document.add_picture(image_stream, width=docx.shared.Cm(4.0))
    except Exception:
        return False
    return True


def job_ad_to_docx_bytes(
    job_ad: JobAdGenerationResult,
    styleguide: str = "",
    *,
    logo_payload: LogoPayload | None = None,
) -> bytes:
    _ = styleguide
    document = docx.Document()
    _add_logo_to_docx(document=document, logo_payload=logo_payload)
    document.add_heading(job_ad.headline or "Stellenanzeige", level=1)
    if has_structured_job_ad_sections(job_ad):
        if job_ad.intro.strip():
            document.add_paragraph(job_ad.intro.strip())
        for heading, items in (
            ("Deine Aufgaben", job_ad.responsibilities),
            ("Dein Profil", job_ad.profile),
            ("Was wir bieten", job_ad.offer),
        ):
            clean_items = dedupe_preserve_order(items)
            if not clean_items:
                continue
            document.add_heading(heading, level=2)
            for item in clean_items:
                document.add_paragraph(item, style="List Bullet")
        if job_ad.cta.strip():
            document.add_paragraph(job_ad.cta.strip())
        if job_ad.equal_opportunity_note.strip():
            document.add_paragraph(job_ad.equal_opportunity_note.strip())
    else:
        document.add_paragraph(build_publishable_job_ad_plain_text(job_ad))
    document.add_heading("Zielgruppe", level=2)
    for item in job_ad.target_group:
        document.add_paragraph(item, style="List Bullet")
    document.add_heading("AGG-Checkliste", level=2)
    for item in job_ad.agg_checklist:
        document.add_paragraph(item, style="List Bullet")
    bio = io.BytesIO()
    document.save(bio)
    return bio.getvalue()


def job_ad_to_pdf_bytes(
    job_ad: JobAdGenerationResult,
    styleguide: str = "",
    *,
    logo_payload: LogoPayload | None = None,
) -> bytes | None:
    _ = styleguide
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import cm
        from reportlab.lib.utils import ImageReader
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import (
            Image,
            ListFlowable,
            ListItem,
            Paragraph,
            SimpleDocTemplate,
            Spacer,
        )
    except Exception:
        return None

    bio = io.BytesIO()
    document = SimpleDocTemplate(
        bio,
        pagesize=A4,
        leftMargin=2 * cm,
        rightMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
        title=job_ad.headline or "Stellenanzeige",
        author="anonymous",
    )
    styles = getSampleStyleSheet()
    story: list[Any] = []

    if logo_payload is not None:
        logo_bytes = logo_payload.get("bytes")
        if isinstance(logo_bytes, bytes) and logo_bytes:
            image_stream = io.BytesIO(logo_bytes)
            try:
                image_width, image_height = ImageReader(image_stream).getSize()
                max_width = 4.2 * cm
                max_height = 1.8 * cm
                scale = min(max_width / image_width, max_height / image_height, 1)
                story.append(
                    Image(
                        io.BytesIO(logo_bytes),
                        width=image_width * scale,
                        height=image_height * scale,
                    )
                )
                story.append(Spacer(1, 0.5 * cm))
            except Exception:
                pass

    def _paragraph(value: str, style_name: str = "BodyText") -> Paragraph:
        return Paragraph(escape(value).replace("\n", "<br/>"), styles[style_name])

    def _append_heading(value: str, style_name: str = "Heading2") -> None:
        if value.strip():
            story.append(_paragraph(value.strip(), style_name))

    def _append_bullets(items: Sequence[str]) -> None:
        clean_items = dedupe_preserve_order(list(items))
        if not clean_items:
            return
        story.append(
            ListFlowable(
                [ListItem(_paragraph(item), leftIndent=0) for item in clean_items],
                bulletType="bullet",
                leftIndent=14,
            )
        )

    _append_heading(job_ad.headline or "Stellenanzeige", "Title")
    if has_structured_job_ad_sections(job_ad):
        if job_ad.intro.strip():
            story.append(_paragraph(job_ad.intro.strip()))
            story.append(Spacer(1, 0.25 * cm))
        for heading, items in (
            ("Deine Aufgaben", job_ad.responsibilities),
            ("Dein Profil", job_ad.profile),
            ("Was wir bieten", job_ad.offer),
        ):
            clean_items = dedupe_preserve_order(items)
            if not clean_items:
                continue
            _append_heading(heading)
            _append_bullets(clean_items)
            story.append(Spacer(1, 0.2 * cm))
        for value in (job_ad.cta, job_ad.equal_opportunity_note):
            if value.strip():
                story.append(_paragraph(value.strip()))
                story.append(Spacer(1, 0.2 * cm))
    else:
        for paragraph in build_publishable_job_ad_plain_text(job_ad).split("\n\n"):
            if paragraph.strip():
                story.append(_paragraph(paragraph.strip()))
                story.append(Spacer(1, 0.2 * cm))

    _append_heading("Zielgruppe")
    _append_bullets(job_ad.target_group)
    _append_heading("AGG-Checkliste")
    _append_bullets(job_ad.agg_checklist)
    document.build(story)
    return bio.getvalue()


def job_ad_preview_shell_options(options: Mapping[str, Any] | None) -> dict[str, Any]:
    raw_options = options if isinstance(options, Mapping) else {}
    tone = str(raw_options.get("tone") or "").strip()
    length = str(raw_options.get("length") or "").strip()
    accent_by_tone = {
        "Professionell & nahbar": "#2563eb",
        "Direkt & pragmatisch": "#374151",
        "Motivierend": "#0f766e",
    }
    return {
        "accent_color": accent_by_tone.get(tone, "#2563eb"),
        "compact": length == "Kompakt",
        "height_px": 620 if length == "Ausführlich" else 560,
    }


def job_ad_preview_html(
    job_ad: JobAdGenerationResult,
    *,
    logo_payload: LogoPayload | None,
) -> str:
    return markdown_article_preview_html(
        build_publishable_job_ad_markdown(job_ad),
        logo_payload=logo_payload,
    )


def estimate_text_area_height(text: str) -> int:
    lines = max(1, len(text.splitlines()))
    return min(520, max(160, 40 + lines * 22))
