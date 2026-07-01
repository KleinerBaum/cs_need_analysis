# parsing.py
"""Input parsing utilities (DOCX/PDF/text) and optional redaction."""

from __future__ import annotations

import io
import re
import unicodedata
import zipfile
from pathlib import Path
from typing import Any, Dict, Tuple
from defusedxml import ElementTree

import docx  # python-docx
import pdfplumber
from constants import (
    SOURCE_UPLOAD_ALLOWED_EXTENSIONS,
    SOURCE_UPLOAD_FILE_SIGNATURES,
    SOURCE_UPLOAD_MAX_BYTES,
    SOURCE_UPLOAD_MAX_BYTES_BY_EXTENSION,
    SOURCE_UPLOAD_TEXT_ENCODINGS,
    SOURCE_UPLOAD_TEXT_MAX_CONTROL_CHAR_RATIO,
)
from docx.table import Table
from docx.text.paragraph import Paragraph


def _normalize_text(text: str) -> str:
    # Normalize newlines and whitespace while preserving structure.
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # Collapse excessive blank lines
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Trim trailing spaces
    text = "\n".join(line.rstrip() for line in text.split("\n"))
    return text.strip()


def read_validated_upload_bytes(upload: Any) -> Tuple[bytes, Dict[str, Any]]:
    """Read upload bytes after cheap size, extension, and signature guards."""

    meta: Dict[str, Any] = {
        "name": getattr(upload, "name", None),
        "type": getattr(upload, "type", None),
        "size": getattr(upload, "size", None),
    }
    extension = _upload_extension(meta.get("name"))
    declared_size = _coerce_upload_size(meta.get("size"))
    if declared_size is not None:
        _validate_upload_size(declared_size, extension)

    try:
        upload.seek(0)
    except (AttributeError, OSError, ValueError):
        pass

    try:
        raw = upload.read()
    except Exception as exc:
        raise ValueError("Datei konnte nicht gelesen werden.") from exc
    finally:
        try:
            upload.seek(0)
        except (AttributeError, OSError, ValueError):
            pass

    if isinstance(raw, bytearray):
        raw = bytes(raw)
    elif isinstance(raw, memoryview):
        raw = raw.tobytes()
    if not isinstance(raw, bytes):
        raise ValueError("Datei konnte nicht gelesen werden.")

    meta["size"] = len(raw)
    _validate_upload_size(len(raw), extension)
    _validate_upload_payload(raw, extension)
    return raw, meta


def _upload_extension(name: Any) -> str:
    extension = Path(str(name or "")).suffix.lower()
    if extension not in SOURCE_UPLOAD_ALLOWED_EXTENSIONS:
        allowed = _format_allowed_extensions()
        raise ValueError(
            f"Dateiformat wird nicht unterstützt. Erlaubt sind {allowed}."
        )
    return extension


def _coerce_upload_size(value: Any) -> int | None:
    if value is None or value == "":
        return None
    try:
        size = int(value)
    except (TypeError, ValueError):
        return None
    return size if size >= 0 else None


def _validate_upload_size(size: int, extension: str) -> None:
    max_bytes = SOURCE_UPLOAD_MAX_BYTES_BY_EXTENSION.get(
        extension, SOURCE_UPLOAD_MAX_BYTES
    )
    if size > max_bytes:
        raise ValueError(
            f"Datei ist zu groß. Maximal erlaubt sind {_format_byte_size(max_bytes)}."
        )


def _validate_upload_payload(raw: bytes, extension: str) -> None:
    if not raw:
        raise ValueError("Datei enthält keinen auslesbaren Inhalt.")

    signatures = SOURCE_UPLOAD_FILE_SIGNATURES.get(extension, ())
    if signatures and not raw.startswith(signatures):
        raise ValueError("Dateisignatur passt nicht zur Dateiendung.")

    if extension == ".docx":
        _validate_docx_container(raw)
    elif extension == ".pdf":
        _validate_pdf_payload(raw)
    elif extension == ".txt":
        _validate_text_payload(raw)


def _validate_docx_container(raw: bytes) -> None:
    try:
        with zipfile.ZipFile(io.BytesIO(raw)) as archive:
            names = set(archive.namelist())
            if "[Content_Types].xml" not in names or "word/document.xml" not in names:
                raise ValueError
    except (zipfile.BadZipFile, ValueError, OSError):
        raise ValueError("DOCX-Datei ist beschädigt oder unvollständig.") from None


def _validate_pdf_payload(raw: bytes) -> None:
    if b"%%EOF" not in raw[-2048:]:
        raise ValueError("PDF-Datei ist beschädigt oder unvollständig.")


def _validate_text_payload(raw: bytes) -> None:
    binary_signatures = tuple(
        signature
        for extension, signatures in SOURCE_UPLOAD_FILE_SIGNATURES.items()
        if extension != ".txt"
        for signature in signatures
    )
    if binary_signatures and raw.startswith(binary_signatures):
        raise ValueError("Dateisignatur passt nicht zur Dateiendung.")
    _decode_text_payload(raw)


def _decode_text_payload(raw: bytes) -> str:
    for encoding in SOURCE_UPLOAD_TEXT_ENCODINGS:
        if encoding == "utf-16" and not raw.startswith((b"\xff\xfe", b"\xfe\xff")):
            continue
        try:
            text = raw.decode(encoding)
        except UnicodeError:
            continue
        if _control_char_ratio(text) <= SOURCE_UPLOAD_TEXT_MAX_CONTROL_CHAR_RATIO:
            return text
    raise ValueError("TXT-Datei wirkt nicht wie eine Textdatei.")


def _control_char_ratio(text: str) -> float:
    if not text:
        return 0.0
    allowed_controls = {"\n", "\r", "\t", "\f"}
    control_count = sum(
        1
        for char in text
        if char not in allowed_controls and unicodedata.category(char) == "Cc"
    )
    return control_count / len(text)


def _format_allowed_extensions() -> str:
    labels = [
        extension.removeprefix(".").upper()
        for extension in SOURCE_UPLOAD_ALLOWED_EXTENSIONS
    ]
    return ", ".join(labels[:-1]) + f" und {labels[-1]}"


def _format_byte_size(size: int) -> str:
    mib = 1024 * 1024
    if size % mib == 0:
        return f"{size // mib} MB"
    return f"{size} Bytes"


def extract_text_from_uploaded_file(upload: Any) -> Tuple[str, Dict[str, Any]]:
    """Extract plain text from a Streamlit UploadedFile.

    Returns:
        (text, meta)
    """
    raw, meta = read_validated_upload_bytes(upload)
    extension = _upload_extension(meta.get("name"))

    pdf_missing_text_layer = False
    if extension == ".docx":
        try:
            text = _extract_docx(raw)
        except Exception as exc:
            raise ValueError(
                "DOCX-Datei ist beschädigt oder unvollständig."
            ) from exc
    elif extension == ".pdf":
        try:
            text, pdf_missing_text_layer = _extract_pdf(raw)
        except Exception as exc:
            raise ValueError("PDF-Datei ist beschädigt oder unvollständig.") from exc
    else:
        text = _decode_text_payload(raw)

    normalized = _normalize_text(text)
    if not normalized:
        if extension == ".docx":
            raise ValueError("DOCX enthält keinen auslesbaren Text.")
        if extension == ".pdf":
            if pdf_missing_text_layer:
                raise ValueError("PDF enthält keinen Textlayer (OCR fehlt).")
            raise ValueError("PDF enthält keinen auslesbaren Text.")
        raise ValueError("Datei enthält keinen auslesbaren Inhalt.")

    return normalized, meta


def extract_text_from_path(path: str | Path) -> str:
    """Helper for local dev: extract text from a local file path."""
    p = Path(path)
    raw = p.read_bytes()
    if p.suffix.lower() == ".docx":
        return _normalize_text(_extract_docx(raw))
    if p.suffix.lower() == ".pdf":
        text, _missing_text_layer = _extract_pdf(raw)
        return _normalize_text(text)
    return _normalize_text(raw.decode("utf-8", errors="replace"))


def _extract_docx(raw: bytes) -> str:
    doc = docx.Document(io.BytesIO(raw))
    out_lines: list[str] = []

    def _append_text(value: str | None) -> None:
        text = (value or "").strip()
        if text:
            out_lines.append(text)

    def _append_table_text(tables: Any) -> None:
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        _append_text(paragraph.text)
                    _append_table_text(cell.tables)

    for paragraph in doc.paragraphs:
        _append_text(paragraph.text)
    _append_table_text(doc.tables)
    for section in doc.sections:
        for part in (section.header, section.footer):
            for paragraph in part.paragraphs:
                _append_text(paragraph.text)
            _append_table_text(part.tables)

    for text in _extract_docx_ooxml_text(raw):
        _append_text(text)

    return "\n".join(dict.fromkeys(out_lines))


def extract_docx_preview_blocks(raw: bytes) -> list[dict[str, Any]]:
    """Return body-order DOCX blocks for an approximate visual preview."""

    _validate_docx_container(raw)
    doc = docx.Document(io.BytesIO(raw))
    blocks: list[dict[str, Any]] = []

    for child in doc.element.body.iterchildren():
        if child.tag.endswith("}p"):
            paragraph = Paragraph(child, doc)
            text = paragraph.text.strip()
            if not text:
                continue
            style_name = str(getattr(paragraph.style, "name", "") or "")
            level = _docx_heading_level(style_name)
            blocks.append(
                {
                    "type": "heading" if level else "paragraph",
                    "text": text,
                    "level": level,
                }
            )
        elif child.tag.endswith("}tbl"):
            table = Table(child, doc)
            rows: list[list[str]] = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    rows.append(cells)
            if rows:
                blocks.append({"type": "table", "rows": rows})

    if not blocks:
        fallback_text = _extract_docx(raw)
        blocks.extend(
            {"type": "paragraph", "text": line, "level": 0}
            for line in fallback_text.splitlines()
            if line.strip()
        )
    return blocks


def _docx_heading_level(style_name: str) -> int:
    normalized = style_name.strip().lower()
    if normalized.startswith("heading"):
        tail = normalized.removeprefix("heading").strip()
        if tail.isdigit():
            return min(max(int(tail), 1), 6)
        return 2
    if normalized.startswith("\u00fcberschrift"):
        tail = normalized.removeprefix("\u00fcberschrift").strip()
        if tail.isdigit():
            return min(max(int(tail), 1), 6)
        return 2
    if normalized in {"title", "titel"}:
        return 1
    return 0


def _extract_docx_ooxml_text(raw: bytes) -> list[str]:
    """Read text nodes that python-docx can miss, e.g. text boxes."""

    lines: list[str] = []
    namespace = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    try:
        with zipfile.ZipFile(io.BytesIO(raw)) as archive:
            xml_names = [
                name
                for name in archive.namelist()
                if name.startswith("word/")
                and name.endswith(".xml")
                and not name.startswith("word/_rels/")
            ]
            for name in xml_names:
                root = ElementTree.fromstring(archive.read(name))
                for text_box in root.iter(f"{namespace}txbxContent"):
                    texts = [
                        node.text.strip()
                        for node in text_box.iter(f"{namespace}t")
                        if node.text and node.text.strip()
                    ]
                    if texts:
                        lines.append(" ".join(texts))
    except (ElementTree.ParseError, KeyError, zipfile.BadZipFile):
        return []
    return lines


def _extract_pdf(raw: bytes) -> tuple[str, bool]:
    out_lines = []
    has_text_layer = False
    with pdfplumber.open(io.BytesIO(raw)) as pdf:
        for page in pdf.pages:
            if getattr(page, "chars", None):
                has_text_layer = True
            t = page.extract_text() or ""
            t = t.strip()
            if t:
                out_lines.append(t)
    return "\n\n".join(out_lines), (not has_text_layer)


_EMAIL_RE = re.compile(r"\b[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}\b", re.IGNORECASE)
_PHONE_RE = re.compile(
    r"(\+?\d{1,3}[\s\-/]?)?(\(?\d{2,5}\)?[\s\-/]?)?\d{3,}[\s\-/]?\d{2,}(?:[\s\-/]?\d{2,})?"
)


def redact_pii(text: str, *, mask: str = "[REDACTED]") -> str:
    """Best-effort redaction to reduce accidental sharing of contact data.

    Notes:
    - This is *not* perfect. It is intended as a safety net.
    - You can disable redaction in the UI if you want to keep contact data.
    """
    # Emails first (high precision)
    text = _EMAIL_RE.sub(mask, text)

    # Phone numbers (lower precision; avoid redacting short numbers)
    def _phone_sub(m: re.Match) -> str:
        s = m.group(0)
        digits = re.sub(r"\D", "", s)
        if len(digits) < 8:
            return s
        return mask

    text = _PHONE_RE.sub(_phone_sub, text)
    return text
