from __future__ import annotations

import json
import base64
import io

import docx
import pytest
from constants import FactKey, SSKey, VACANCY_DRAFT_SCHEMA_VERSION
from document_preview import (
    article_preview_html,
    document_preview_shell,
    markdown_article_preview_html,
    pdf_preview_html,
)
from schemas import (
    BooleanSearchChannelQueries,
    BooleanSearchPack,
    InterviewPrepSheetHiringManager,
    JobAdExtract,
    VacancyBrief,
    VacancyStructuredData,
)
from summary_exports import (
    boolean_search_pack_to_markdown,
    brief_to_markdown,
    build_live_artifact_preview_payload,
    build_summary_input_fingerprint,
    build_vacancy_draft_payload,
    parse_vacancy_draft_json,
    vacancy_draft_payload_to_json,
)
from wizard_pages.summary_exporters import (
    _brief_to_docx_bytes,
    _interview_prep_fach_to_pdf_bytes,
)


def _fingerprint(*, answers: dict[str, object]) -> str:
    return build_summary_input_fingerprint(
        job=JobAdExtract(
            job_title="Data Engineer",
            company_name="ACME GmbH",
            location_country="DE",
        ),
        answers=answers,
        selected_role_tasks=["Build data products"],
        selected_skills=["Python", "SQL"],
        selected_benefits=["Mentoring"],
        esco_occupation_selected={"uri": "uri:occ:1", "title": "Data Engineer"},
        esco_match_explainability={
            "reason": "Anker bestätigt",
            "confidence": "high",
            "provenance": ["exact label match"],
        },
        esco_selected_skills_must=[{"uri": "uri:skill:python", "title": "Python"}],
        esco_selected_skills_nice=[],
    )


def test_build_summary_input_fingerprint_is_stable_and_sensitive_to_inputs() -> None:
    baseline = _fingerprint(answers={"team_size": 5})

    assert baseline == _fingerprint(answers={"team_size": 5})
    assert baseline != _fingerprint(answers={"team_size": 10})
    assert len(baseline) == 64


def test_brief_to_markdown_exports_primary_sections() -> None:
    brief = VacancyBrief(
        one_liner="One line",
        hiring_context="Context",
        role_summary="Summary",
        top_responsibilities=["Build pipelines"],
        must_have=["Python"],
        nice_to_have=["SQL"],
        dealbreakers=["No cloud experience"],
        interview_plan=["HR screen"],
        evaluation_rubric=["Technical depth"],
        risks_open_questions=["Budget open"],
        job_ad_draft="Draft text",
        structured_data=VacancyStructuredData(
            job_extract={"job_title": "Data Engineer"},
            answers={},
        ),
    )

    markdown = brief_to_markdown(brief)

    assert markdown.startswith("# Recruiting Brief - Data Engineer")
    assert "## Wichtigste Aufgaben\n- Build pipelines" in markdown
    assert "## Stellenanzeigenentwurf (DE)\nDraft text" in markdown

    english_markdown = brief_to_markdown(brief, language="en")

    assert english_markdown.startswith("# Recruiting brief - Data Engineer")
    assert "## Top responsibilities\n- Build pipelines" in english_markdown
    assert "## Job ad draft\nDraft text" in english_markdown


def test_brief_to_markdown_exports_salary_caveat_by_language() -> None:
    brief = VacancyBrief(
        one_liner="One line",
        hiring_context="Context",
        role_summary="Summary",
        top_responsibilities=[],
        must_have=[],
        nice_to_have=[],
        dealbreakers=[],
        interview_plan=[],
        evaluation_rubric=[],
        risks_open_questions=[],
        job_ad_draft="Draft text",
        structured_data=VacancyStructuredData(
            job_extract={"job_title": "Data Engineer"},
            answers={},
            offer_positioning={
                "salary_caveat": "Salary forecast is orientation only."
            },
        ),
    )

    markdown = brief_to_markdown(brief)
    english_markdown = brief_to_markdown(brief, language="en")

    assert "## Vergütungshinweis\n- Salary forecast is orientation only." in markdown
    assert "## Compensation note\n- Salary forecast is orientation only." in english_markdown


def test_brief_to_docx_exports_headings_by_language() -> None:
    brief = VacancyBrief(
        one_liner="One line",
        hiring_context="Context",
        role_summary="Summary",
        top_responsibilities=["Build pipelines"],
        must_have=["Python"],
        nice_to_have=["SQL"],
        dealbreakers=[],
        interview_plan=["HR screen"],
        evaluation_rubric=[],
        risks_open_questions=["Budget open"],
        job_ad_draft="Draft text",
        structured_data=VacancyStructuredData(
            job_extract={"job_title": "Data Engineer"},
            answers={},
        ),
    )

    german_document = docx.Document(io.BytesIO(_brief_to_docx_bytes(brief)))
    german_text = "\n".join(paragraph.text for paragraph in german_document.paragraphs)
    english_document = docx.Document(
        io.BytesIO(_brief_to_docx_bytes(brief, language="en"))
    )
    english_text = "\n".join(paragraph.text for paragraph in english_document.paragraphs)

    assert "Recruiting Brief - Data Engineer" in german_text
    assert "Wichtigste Aufgaben" in german_text
    assert "Stellenanzeigenentwurf (DE)" in german_text
    assert "Recruiting brief - Data Engineer" in english_text
    assert "Top responsibilities" in english_text
    assert "Job ad draft" in english_text


def test_interview_fach_pdf_uses_english_heading_when_reportlab_available() -> None:
    pytest.importorskip("reportlab")
    sheet = InterviewPrepSheetHiringManager(
        role_title="Senior Data Engineer",
        interview_stage="Technical interview",
        duration_minutes=60,
        competencies_to_validate=["Python"],
        question_blocks=[],
        technical_deep_dive_topics=[],
        case_or_task_prompt=None,
        evaluation_rubric=[],
        hiring_signal_summary=[],
        debrief_questions=["Debrief"],
    )

    pdf_bytes = _interview_prep_fach_to_pdf_bytes(sheet, language="en")

    assert pdf_bytes is not None
    assert b"Hiring manager sheet" in pdf_bytes
    assert b"Fachbereich-Sheet" not in pdf_bytes


def test_boolean_search_pack_to_markdown_formats_channel_queries() -> None:
    channel = BooleanSearchChannelQueries(
        broad=["site:linkedin.com/in Python"],
        focused=[],
        fallback=["Data Engineer Berlin"],
    )
    pack = BooleanSearchPack(
        role_title="Data Engineer",
        target_locations=["Berlin"],
        seniority_terms=["Senior"],
        must_have_terms=["Python"],
        exclusion_terms=[],
        google=channel,
        linkedin=channel,
        xing=channel,
        channel_limitations=["LinkedIn search syntax varies"],
        usage_notes=["Start broad, then tighten"],
    )

    markdown = boolean_search_pack_to_markdown(pack)

    assert markdown.startswith("# Suchstrings")
    assert "**Rolle:** Data Engineer" in markdown
    assert "- `site:linkedin.com/in Python`" in markdown
    assert "### Fokussiert\n- —" in markdown
    assert "## Nutzungshinweise\n- Start broad, then tighten" in markdown

    english_markdown = boolean_search_pack_to_markdown(pack, language="en")

    assert english_markdown.startswith("# Search strings")
    assert "**Role title:** Data Engineer" in english_markdown
    assert "### Focused\n- —" in english_markdown
    assert "## Usage notes\n- Start broad, then tighten" in english_markdown


def test_live_artifact_preview_payload_uses_current_inputs_without_contact_details() -> None:
    payload = build_live_artifact_preview_payload(
        job=JobAdExtract(
            job_title="Data Engineer",
            company_name="ACME GmbH",
            responsibilities=["Build data products"],
            must_have_skills=["Python"],
            benefits=["Hybrid work"],
            location_city="Berlin",
        ),
        answers={},
        selected_role_tasks=["Own data pipelines"],
        selected_skills=["SQL"],
        selected_benefits=["Mentoring"],
        interview_process={
            "candidate_stages": ["HR Screen", "Technical interview"],
            "selected_values": [
                {
                    "Bereich": "Interne Rollen",
                    "Feld": "Money E-Mail-Adresse",
                    "Wert": "person@example.com",
                },
                {
                    "Bereich": "Interview",
                    "Feld": "Interviewphase 1",
                    "Wert": "Technical interview",
                },
            ],
        },
    )

    assert payload["is_preview"] is True
    assert payload["fragments"]["brief"]["summary"] == "Data Engineer bei ACME GmbH"
    assert any(
        "Own data pipelines" in bullet
        for bullet in payload["fragments"]["job_ad"]["bullets"]
    )
    assert any(
        "SQL" in bullet for bullet in payload["fragments"]["boolean_search"]["bullets"]
    )
    interview_bullets = "\n".join(payload["fragments"]["interview_hr"]["bullets"])
    fach_bullets = "\n".join(payload["fragments"]["interview_fach"]["bullets"])
    assert "Technical interview" in interview_bullets
    assert "Technical interview" in fach_bullets
    assert "person@example.com" not in interview_bullets
    assert "person@example.com" not in fach_bullets


def test_live_artifact_preview_payload_uses_english_copy() -> None:
    payload = build_live_artifact_preview_payload(
        job=JobAdExtract(
            job_title="Data Engineer",
            company_name="ACME GmbH",
            responsibilities=["Build data products"],
            must_have_skills=["Python"],
            benefits=["Hybrid work"],
            location_city="Berlin",
            salary_range={"min": 70000, "max": 90000, "currency": "EUR"},
        ),
        answers={},
        selected_role_tasks=["Own data pipelines"],
        selected_skills=["SQL"],
        selected_benefits=["Mentoring"],
        intake_facts={},
        interview_process={"candidate_stages": ["HR Screen"]},
        language="en",
    )

    assert payload["notice"].startswith("Live preview from current inputs")
    assert payload["fragments"]["brief"]["title"] == "Recruiting brief"
    assert payload["fragments"]["brief"]["summary"] == "Data Engineer at ACME GmbH"
    assert payload["context"]["salary"] == "70000.0 - 90000.0 EUR"
    assert "Offer: Mentoring, Hybrid work" in payload["fragments"]["brief"]["bullets"]
    assert "Recruiting-Unterlagen" not in json.dumps(payload, ensure_ascii=False)


def test_live_artifact_preview_payload_uses_role_outcome_facts() -> None:
    payload = build_live_artifact_preview_payload(
        job=JobAdExtract(
            job_title="Product Lead",
            responsibilities=["Coordinate delivery"],
            must_have_skills=["Roadmapping"],
        ),
        selected_role_tasks=[],
        selected_skills=["Stakeholder management"],
        intake_facts={
            FactKey.ROLE_BUSINESS_OUTCOME_PRIMARY.value: "Reduce release cycle time",
            FactKey.ROLE_DELIVERABLES.value: ["Release dashboard"],
            FactKey.ROLE_SUCCESS_METRICS_TIMELINE.value: {
                "90_days": "Dashboard live with weekly adoption review",
            },
            FactKey.ROLE_DECISION_SCOPE.value: "fachliche_empfehlungen",
            FactKey.ROLE_RESPONSIBILITIES_PRIORITIZED.value: [
                {"label": "Own delivery roadmap", "priority": "must"},
            ],
            FactKey.COMPANY_NON_NEGOTIABLES.value: ["Berlin"],
        },
    )

    brief_bullets = "\n".join(payload["fragments"]["brief"]["bullets"])
    job_ad_bullets = "\n".join(payload["fragments"]["job_ad"]["bullets"])
    search_bullets = "\n".join(payload["fragments"]["boolean_search"]["bullets"])
    interview_bullets = "\n".join(payload["fragments"]["interview_hr"]["bullets"])
    fach_bullets = "\n".join(payload["fragments"]["interview_fach"]["bullets"])

    assert "Reduce release cycle time" in brief_bullets
    assert "Release dashboard" in job_ad_bullets
    assert "Nicht verhandelbar: Berlin" in search_bullets
    assert "Dashboard live with weekly adoption review" in interview_bullets
    assert "Dashboard live with weekly adoption review" in fach_bullets
    assert payload["context"]["output_count"] == 1
    assert set(payload["fragments"]) == {
        "brief",
        "job_ad",
        "boolean_search",
        "interview_hr",
        "interview_fach",
    }


def test_vacancy_draft_payload_exports_schema_and_allowed_state_only() -> None:
    payload = build_vacancy_draft_payload(
        {
            SSKey.SOURCE_TEXT.value: "Synthetic jobspec",
            SSKey.ANSWERS.value: {"company_name": "Example GmbH"},
            SSKey.MODEL.value: "runtime-model",
            SSKey.USAGE_EVENTS.value: [{"event_type": "artifact_generated"}],
            "unknown": "ignored",
        },
        allowed_keys=(SSKey.SOURCE_TEXT, SSKey.ANSWERS),
    )

    assert payload["schema"] == "cs_need_analysis.vacancy_draft"
    assert payload["schema_version"] == VACANCY_DRAFT_SCHEMA_VERSION
    assert payload["state"] == {
        SSKey.SOURCE_TEXT.value: "Synthetic jobspec",
        SSKey.ANSWERS.value: {"company_name": "Example GmbH"},
    }

    parsed = parse_vacancy_draft_json(vacancy_draft_payload_to_json(payload))

    assert parsed == payload


def test_markdown_article_preview_html_escapes_hostile_markdown() -> None:
    html = markdown_article_preview_html(
        "# <script>alert(1)</script>\n\n"
        'Intro <img src=x onerror="alert(1)">\n\n'
        '## <svg onload="alert(1)"></svg>\n\n'
        '- <a href="javascript:alert(1)">bad</a>'
    )

    assert "<script>" not in html
    assert "<img" not in html
    assert "<svg" not in html
    assert "<a href=" not in html
    assert "&lt;script&gt;alert(1)&lt;/script&gt;" in html
    assert "&lt;img src=x onerror=&quot;alert(1)&quot;&gt;" in html
    assert "&lt;svg onload=&quot;alert(1)&quot;&gt;&lt;/svg&gt;" in html
    assert "&lt;a href=&quot;javascript:alert(1)&quot;&gt;bad&lt;/a&gt;" in html


def test_article_preview_html_escapes_hostile_structured_fields() -> None:
    html = article_preview_html(
        headline='<script>alert("headline")</script>',
        intro='<img src=x onerror="alert(1)">',
        sections=[
            (
                '<svg onload="alert(1)"></svg>',
                ['<a href="javascript:alert(1)">bad</a>'],
            )
        ],
        closing=['<iframe src="javascript:alert(1)"></iframe>'],
    )

    assert "<script>" not in html
    assert "<img" not in html
    assert "<svg" not in html
    assert "<a href=" not in html
    assert "<iframe" not in html
    assert "&lt;script&gt;alert(&quot;headline&quot;)&lt;/script&gt;" in html
    assert "&lt;img src=x onerror=&quot;alert(1)&quot;&gt;" in html
    assert "&lt;svg onload=&quot;alert(1)&quot;&gt;&lt;/svg&gt;" in html
    assert "&lt;a href=&quot;javascript:alert(1)&quot;&gt;bad&lt;/a&gt;" in html
    assert "&lt;iframe src=&quot;javascript:alert(1)&quot;&gt;&lt;/iframe&gt;" in html


def test_document_preview_shell_escapes_title_and_rejects_css_injection() -> None:
    html = document_preview_shell(
        "<article>trusted</article>",
        title="<script>alert(1)</script>",
        accent_color='red; background-image: url("javascript:alert(1)")',
    )

    assert "<script>" not in html
    assert "&lt;script&gt;alert(1)&lt;/script&gt;" in html
    assert "javascript:alert" not in html
    assert 'background-image: url("javascript:alert(1)")' not in html
    assert "border-top: 5px solid #2563eb;" in html


def test_pdf_preview_html_uses_base64_data_uri_without_raw_payload() -> None:
    html = pdf_preview_html(b"%PDF-1.4\n<script>alert(1)</script>\n%%EOF\n")

    assert html.startswith(
        '<iframe class="cs-document-pdf-frame" title="Dokumentvorschau" '
        'src="data:application/pdf;base64,'
    )
    assert "<script>" not in html
    assert "%%EOF" not in html


def test_markdown_article_preview_html_accepts_normalized_logo_payload() -> None:
    png_bytes = base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO5vS3wAAAAASUVORK5CYII="
    )
    html = markdown_article_preview_html(
        "# Data Engineer",
        logo_payload={
            "name": "brand.png",
            "mime_type": "image/png",
            "extension": ".png",
            "byte_size": len(png_bytes),
            "width_px": 1,
            "height_px": 1,
            "bytes": png_bytes,
        },
    )

    assert "data:image/png;base64," in html
    assert base64.b64encode(png_bytes).decode("ascii") in html
