from __future__ import annotations

import io
from zipfile import ZipFile

import docx
import exporters.job_ad as job_ad_exporter
import summary_job_ad
from document_preview import markdown_article_preview_html
from llm_client import JobAdGenerationResult
from summary_job_ad import (
    build_publishable_job_ad_markdown,
    build_publishable_job_ad_plain_text,
    dedupe_preserve_order,
    estimate_text_area_height,
    job_ad_to_docx_bytes,
    normalize_list_item,
    sanitize_generated_job_ad,
)


def test_summary_job_ad_facade_reexports_job_ad_exporter_helpers() -> None:
    assert summary_job_ad.job_ad_to_docx_bytes is job_ad_exporter.job_ad_to_docx_bytes
    assert summary_job_ad.job_ad_to_pdf_bytes is job_ad_exporter.job_ad_to_pdf_bytes
    assert summary_job_ad.job_ad_preview_html is job_ad_exporter.job_ad_preview_html
    assert (
        summary_job_ad.build_publishable_job_ad_plain_text
        is job_ad_exporter.build_publishable_job_ad_plain_text
    )


def test_normalize_list_item_removes_common_bullet_prefixes() -> None:
    assert normalize_list_item("- Data Engineer") == "Data Engineer"
    assert normalize_list_item("1. Python") == "Python"
    assert normalize_list_item("• SQL") == "SQL"


def test_dedupe_preserve_order_strips_empty_and_case_duplicates() -> None:
    assert dedupe_preserve_order([" Python ", "", "python", "SQL"]) == [
        "Python",
        "SQL",
    ]


def test_sanitize_generated_job_ad_removes_embedded_sections() -> None:
    source = JobAdGenerationResult(
        headline=" Baggerfahrer (m/w/d) ",
        target_group=["Baggerfahrer/innen"],
        agg_checklist=["Geschlechtsneutrale Ansprache vorhanden."],
        job_ad_text=(
            "Wir suchen Verstärkung für unser Team.\n\n"
            "Hinweis: Startdatum ist nicht angegeben.\n"
            "Zielgruppe\n"
            "- Galabauer/innen\n"
            "AGG-Checkliste\n"
            "- Fehlende Angaben: Bewerbungsschluss nicht angegeben.\n"
        ),
    )

    sanitized, notes = sanitize_generated_job_ad(source)

    assert sanitized.headline == "Baggerfahrer (m/w/d)"
    assert sanitized.job_ad_text == "Wir suchen Verstärkung für unser Team."
    assert sanitized.target_group == ["Baggerfahrer/innen", "Galabauer/innen"]
    assert (
        "Fehlende Angaben: Bewerbungsschluss nicht angegeben."
        in sanitized.agg_checklist
    )
    assert notes == ["Startdatum ist nicht angegeben."]


def test_job_ad_generation_result_accepts_structured_sections_and_legacy_payload() -> None:
    legacy = JobAdGenerationResult.model_validate(
        {
            "headline": "Titel",
            "target_group": [],
            "agg_checklist": [],
            "job_ad_text": "Legacy text",
        }
    )
    structured = JobAdGenerationResult.model_validate(
        {
            "headline": "Titel",
            "target_group": ["Senior Professionals"],
            "agg_checklist": ["AGG-konform formuliert."],
            "job_ad_text": "Fallback",
            "intro": "Intro",
            "responsibilities": ["Aufgabe"],
            "profile": ["Profil"],
            "offer": ["Benefit"],
            "cta": "Jetzt bewerben.",
            "equal_opportunity_note": "Alle Bewerbungen sind willkommen.",
        }
    )

    assert legacy.intro == ""
    assert structured.responsibilities == ["Aufgabe"]


def test_publishable_job_ad_markdown_uses_structured_sections_without_raw_markdown() -> None:
    source = JobAdGenerationResult(
        headline="**Senior Engineer**",
        target_group=[],
        agg_checklist=[],
        job_ad_text="Fallback",
        intro="Gestalte Plattformen.",
        responsibilities=["**Baue Pipelines**"],
        profile=["Python"],
        offer=["Mentoring"],
        cta="Bewirb dich jetzt.",
        equal_opportunity_note="Alle Menschen sind willkommen.",
    )
    sanitized, _ = sanitize_generated_job_ad(source)

    markdown = build_publishable_job_ad_markdown(sanitized)
    plain_text = build_publishable_job_ad_plain_text(sanitized)

    assert "# Senior Engineer" in markdown
    assert "- Baue Pipelines" in markdown
    assert "**" not in markdown
    assert "# " not in plain_text


def test_publishable_job_ad_markdown_uses_language_specific_static_headings() -> None:
    job_ad = JobAdGenerationResult(
        headline="Senior Engineer",
        target_group=[],
        agg_checklist=[],
        job_ad_text="Fallback",
        intro="Gestalte Plattformen.",
        responsibilities=["Baue Pipelines"],
        profile=["Python"],
        offer=["Mentoring"],
    )

    german_markdown = build_publishable_job_ad_markdown(job_ad, language="de")
    english_markdown = build_publishable_job_ad_markdown(job_ad, language="en")

    assert "## Deine Aufgaben\n\n- Baue Pipelines" in german_markdown
    assert "## Dein Profil\n\n- Python" in german_markdown
    assert "## Was wir bieten\n\n- Mentoring" in german_markdown
    assert "## Your responsibilities\n\n- Baue Pipelines" in english_markdown
    assert "## Your profile\n\n- Python" in english_markdown
    assert "## What we offer\n\n- Mentoring" in english_markdown
    assert "Gestalte Plattformen." in english_markdown
    assert "Deine Aufgaben" not in english_markdown


def test_job_ad_to_docx_bytes_exports_publishable_sections_without_styleguide() -> None:
    job_ad = JobAdGenerationResult(
        headline="Senior Engineer",
        target_group=["Senior Professionals"],
        agg_checklist=["AGG-konform formuliert."],
        job_ad_text="Fallback",
        intro="Gestalte Plattformen.",
        responsibilities=["Baue Pipelines"],
        profile=["Python"],
        offer=["Mentoring"],
        cta="Bewirb dich jetzt.",
        equal_opportunity_note="Alle Menschen sind willkommen.",
    )

    docx_bytes = job_ad_to_docx_bytes(
        job_ad,
        styleguide="Styleguide darf nicht exportiert werden.",
    )

    with ZipFile(io.BytesIO(docx_bytes), "r") as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")
    assert "Senior Engineer" in document_xml
    assert "Baue Pipelines" in document_xml
    assert "Styleguide darf nicht exportiert werden" not in document_xml


def test_job_ad_to_docx_bytes_uses_language_specific_static_headings() -> None:
    job_ad = JobAdGenerationResult(
        headline="Senior Engineer",
        target_group=["Senior Professionals"],
        agg_checklist=["AGG-konform formuliert."],
        job_ad_text="Fallback",
        intro="Gestalte Plattformen.",
        responsibilities=["Baue Pipelines"],
        profile=["Python"],
        offer=["Mentoring"],
    )

    german_document = docx.Document(
        io.BytesIO(job_ad_to_docx_bytes(job_ad, language="de"))
    )
    german_text = "\n".join(paragraph.text for paragraph in german_document.paragraphs)
    english_document = docx.Document(
        io.BytesIO(job_ad_to_docx_bytes(job_ad, language="en"))
    )
    english_text = "\n".join(
        paragraph.text for paragraph in english_document.paragraphs
    )

    assert "Deine Aufgaben" in german_text
    assert "Dein Profil" in german_text
    assert "Was wir bieten" in german_text
    assert "Zielgruppe" in german_text
    assert "AGG-Checkliste" in german_text
    assert "Your responsibilities" in english_text
    assert "Your profile" in english_text
    assert "What we offer" in english_text
    assert "Target group" in english_text
    assert "AGG checklist" in english_text
    assert "Baue Pipelines" in english_text
    assert "Deine Aufgaben" not in english_text


def test_job_ad_to_docx_bytes_uses_language_specific_fallback_title() -> None:
    job_ad = JobAdGenerationResult(
        headline="",
        target_group=[],
        agg_checklist=[],
        job_ad_text="Fallback",
    )

    german_document = docx.Document(
        io.BytesIO(job_ad_to_docx_bytes(job_ad, language="de"))
    )
    english_document = docx.Document(
        io.BytesIO(job_ad_to_docx_bytes(job_ad, language="en"))
    )

    assert german_document.paragraphs[0].text == "Stellenanzeige"
    assert english_document.paragraphs[0].text == "Job ad"


def test_markdown_article_preview_html_renders_publishable_job_ad_text() -> None:
    preview_html = markdown_article_preview_html(
        "# Senior Engineer\n\n"
        "Gestalte Plattformen.\n\n"
        "## Deine Aufgaben\n\n"
        "- Baue Pipelines\n"
        "- Sichere Datenqualitaet"
    )

    assert "<h1>Senior Engineer</h1>" in preview_html
    assert "<p>Gestalte Plattformen.</p>" in preview_html
    assert "<h2>Deine Aufgaben</h2>" in preview_html
    assert "<li>Baue Pipelines</li>" in preview_html
    assert "<li>Sichere Datenqualitaet</li>" in preview_html


def test_markdown_article_preview_html_splits_long_job_ad_into_two_pages() -> None:
    markdown = (
        "# Senior Engineer\n\n"
        "Gestalte Plattformen.\n\n"
        "## Deine Aufgaben\n\n"
        + "\n".join(f"- Aufgabe {index}" for index in range(1, 18))
        + "\n\n## Dein Profil\n\n- Python\n- Datenqualitaet"
    )

    preview_html = markdown_article_preview_html(markdown)

    assert 'class="cs-document-spread"' in preview_html
    assert preview_html.count('class="cs-document-page"') == 2
    assert "<li>Aufgabe 1</li>" in preview_html
    assert "<li>Datenqualitaet</li>" in preview_html


def test_estimate_text_area_height_has_minimum_and_cap() -> None:
    assert estimate_text_area_height("Kurz") == 160
    assert estimate_text_area_height("\n".join(["Zeile"] * 50)) == 520
