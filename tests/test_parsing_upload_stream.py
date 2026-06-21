from __future__ import annotations

import io
import zipfile

import document_preview
import docx
import pytest

from constants import SOURCE_UPLOAD_MAX_BYTES
import parsing
from parsing import extract_docx_preview_blocks, extract_text_from_uploaded_file


class _FakeUpload:
    def __init__(
        self,
        payload: bytes,
        *,
        name: str = "jobspec.txt",
        size: int | None = None,
    ) -> None:
        self.name = name
        self.type = "text/plain"
        self.size = len(payload) if size is None else size
        self._stream = io.BytesIO(payload)
        self.read_calls = 0

    def seek(self, pos: int) -> int:
        return self._stream.seek(pos)

    def read(self) -> bytes:
        self.read_calls += 1
        return self._stream.read()


def _docx_payload(build_fn) -> bytes:
    document = docx.Document()
    build_fn(document)
    out = io.BytesIO()
    document.save(out)
    return out.getvalue()


def _pdf_payload(text: str) -> bytes:
    from reportlab.pdfgen import canvas

    out = io.BytesIO()
    pdf = canvas.Canvas(out)
    pdf.drawString(72, 720, text)
    pdf.save()
    return out.getvalue()


def _malformed_docx_container() -> bytes:
    out = io.BytesIO()
    with zipfile.ZipFile(out, "w") as archive:
        archive.writestr(
            "[Content_Types].xml",
            (
                '<?xml version="1.0" encoding="UTF-8"?>'
                '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
                '<Default Extension="rels" '
                'ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
                '<Default Extension="xml" ContentType="application/xml"/>'
                '<Override PartName="/word/document.xml" '
                'ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
                "</Types>"
            ),
        )
        archive.writestr("word/document.xml", "<w:document>")
    return out.getvalue()


def test_extract_text_from_uploaded_file_is_stable_across_multiple_reads() -> None:
    upload = _FakeUpload(b"Senior Data Engineer\nPython")

    first_text, first_meta = extract_text_from_uploaded_file(upload)
    second_text, second_meta = extract_text_from_uploaded_file(upload)

    assert first_text == "Senior Data Engineer\nPython"
    assert second_text == first_text
    assert first_meta == second_meta


def test_extract_text_from_uploaded_file_raises_on_empty_payload() -> None:
    upload = _FakeUpload(b"")

    with pytest.raises(ValueError, match="Datei enthält keinen auslesbaren Inhalt"):
        extract_text_from_uploaded_file(upload)


def test_extract_text_from_uploaded_file_reads_docx_tables() -> None:
    buf = io.BytesIO()
    doc = docx.Document()
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Nur"
    table.rows[0].cells[1].text = "Tabelle"
    doc.save(buf)

    upload = _FakeUpload(
        buf.getvalue(),
        name="jobspec.docx",
    )

    text, _meta = extract_text_from_uploaded_file(upload)

    assert text == "Nur\nTabelle"


def test_extract_text_from_uploaded_file_reads_docx_headers_and_footers() -> None:
    def _build(document: docx.Document) -> None:
        document.sections[0].header.paragraphs[0].text = "Senior AI Consultant"
        document.add_paragraph("Main body requirements")
        document.sections[0].footer.paragraphs[0].text = "Reference R001"

    upload = _FakeUpload(_docx_payload(_build), name="jobspec.docx")

    text, _meta = extract_text_from_uploaded_file(upload)

    assert "Senior AI Consultant" in text
    assert "Main body requirements" in text
    assert "Reference R001" in text


def test_extract_text_from_uploaded_file_reads_small_valid_pdf() -> None:
    upload = _FakeUpload(
        _pdf_payload("Senior Data Engineer"),
        name="jobspec.pdf",
    )

    text, _meta = extract_text_from_uploaded_file(upload)

    assert "Senior Data Engineer" in text


def test_extract_text_from_uploaded_file_rejects_declared_oversize_before_read() -> None:
    upload = _FakeUpload(
        b"Senior Data Engineer",
        name="jobspec.txt",
        size=SOURCE_UPLOAD_MAX_BYTES + 1,
    )

    with pytest.raises(ValueError, match="Datei ist zu groß"):
        extract_text_from_uploaded_file(upload)

    assert upload.read_calls == 0


def test_extract_text_from_uploaded_file_rejects_actual_oversize_after_read() -> None:
    upload = _FakeUpload(
        b"x" * (SOURCE_UPLOAD_MAX_BYTES + 1),
        name="jobspec.txt",
        size=1,
    )

    with pytest.raises(ValueError, match="Datei ist zu groß"):
        extract_text_from_uploaded_file(upload)

    assert upload.read_calls == 1


def test_extract_text_from_uploaded_file_rejects_unsupported_extension_before_read() -> None:
    upload = _FakeUpload(b"{\\rtf1 jobspec}", name="jobspec.rtf")

    with pytest.raises(ValueError, match="Dateiformat wird nicht unterstützt"):
        extract_text_from_uploaded_file(upload)

    assert upload.read_calls == 0


def test_extract_text_from_uploaded_file_rejects_wrong_extension_signature(
    monkeypatch,
) -> None:
    upload = _FakeUpload(b"%PDF-1.4\n%%EOF\n", name="jobspec.docx")
    monkeypatch.setattr(
        parsing,
        "_extract_docx",
        lambda _raw: pytest.fail("DOCX parser must not be called"),
    )

    with pytest.raises(ValueError, match="Dateisignatur passt nicht"):
        extract_text_from_uploaded_file(upload)


def test_extract_text_from_uploaded_file_rejects_corrupt_docx_before_parser(
    monkeypatch,
) -> None:
    upload = _FakeUpload(b"PK\x03\x04not-a-valid-zip", name="jobspec.docx")
    monkeypatch.setattr(
        parsing,
        "_extract_docx",
        lambda _raw: pytest.fail("DOCX parser must not be called"),
    )

    with pytest.raises(ValueError, match="DOCX-Datei ist beschädigt"):
        extract_text_from_uploaded_file(upload)


def test_extract_text_from_uploaded_file_maps_malformed_docx_parser_error() -> None:
    upload = _FakeUpload(_malformed_docx_container(), name="jobspec.docx")

    with pytest.raises(ValueError) as exc_info:
        extract_text_from_uploaded_file(upload)

    message = str(exc_info.value)
    assert message == "DOCX-Datei ist beschädigt oder unvollständig."
    assert "Package" not in message
    assert "XML" not in message


def test_extract_text_from_uploaded_file_rejects_corrupt_pdf_before_parser(
    monkeypatch,
) -> None:
    upload = _FakeUpload(b"%PDF-1.4\nnot complete", name="jobspec.pdf")
    monkeypatch.setattr(
        parsing,
        "_extract_pdf",
        lambda _raw: pytest.fail("PDF parser must not be called"),
    )

    with pytest.raises(ValueError, match="PDF-Datei ist beschädigt"):
        extract_text_from_uploaded_file(upload)


def test_extract_text_from_uploaded_file_maps_pdf_parser_error(monkeypatch) -> None:
    upload = _FakeUpload(b"%PDF-1.4\n%%EOF\n", name="jobspec.pdf")

    def fake_extract_pdf(_raw: bytes) -> tuple[str, bool]:
        raise RuntimeError("private parser detail")

    monkeypatch.setattr(parsing, "_extract_pdf", fake_extract_pdf)

    with pytest.raises(ValueError) as exc_info:
        extract_text_from_uploaded_file(upload)

    message = str(exc_info.value)
    assert message == "PDF-Datei ist beschädigt oder unvollständig."
    assert "private parser detail" not in message


def test_extract_text_from_uploaded_file_rejects_binary_txt() -> None:
    upload = _FakeUpload(b"\x00\x01\x02\x03", name="jobspec.txt")

    with pytest.raises(ValueError, match="TXT-Datei wirkt nicht wie eine Textdatei"):
        extract_text_from_uploaded_file(upload)


def test_extract_text_from_uploaded_file_rejects_control_heavy_utf16_txt() -> None:
    upload = _FakeUpload(b"\xff\xfe" + (b"\x00\x00" * 64), name="jobspec.txt")

    with pytest.raises(ValueError, match="TXT-Datei wirkt nicht wie eine Textdatei"):
        extract_text_from_uploaded_file(upload)


def test_extract_docx_preview_blocks_preserves_body_order() -> None:
    def _build(document: docx.Document) -> None:
        document.add_heading("Senior Data Engineer", level=1)
        document.add_paragraph("Build reliable data products.")
        table = document.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "Python"
        table.rows[0].cells[1].text = "Airflow"

    blocks = extract_docx_preview_blocks(_docx_payload(_build))

    assert blocks[0] == {
        "type": "heading",
        "text": "Senior Data Engineer",
        "level": 1,
    }
    assert blocks[1] == {
        "type": "paragraph",
        "text": "Build reliable data products.",
        "level": 0,
    }
    assert blocks[2] == {"type": "table", "rows": [["Python", "Airflow"]]}


def test_extract_text_from_uploaded_file_pdf_without_ocr_has_specific_error(
    monkeypatch,
) -> None:
    upload = _FakeUpload(b"%PDF-1.4\n%%EOF\n", name="scan.pdf")

    monkeypatch.setattr(parsing, "_extract_pdf", lambda _raw: ("", True))

    with pytest.raises(
        ValueError, match="PDF enthält keinen Textlayer \\(OCR fehlt\\)"
    ):
        extract_text_from_uploaded_file(upload)


def test_text_preview_html_escapes_hostile_html() -> None:
    html = document_preview.text_preview_html(
        'Hello <img src=x onerror="alert(1)">\n<script>alert(1)</script>'
    )

    assert "<img" not in html
    assert "<script>" not in html
    assert "&lt;img src=x onerror=&quot;alert(1)&quot;&gt;" in html
    assert "&lt;script&gt;alert(1)&lt;/script&gt;" in html


def test_docx_preview_html_escapes_hostile_paragraphs_and_table_cells(
    monkeypatch,
) -> None:
    monkeypatch.setattr(
        document_preview,
        "extract_docx_preview_blocks",
        lambda _raw: [
            {"type": "paragraph", "text": '<img src=x onerror="alert(1)">', "level": 0},
            {"type": "table", "rows": [["<script>alert(1)</script>"]]},
        ],
    )

    html = document_preview.docx_preview_html(b"docx")

    assert "<img" not in html
    assert "<script>" not in html
    assert "&lt;img src=x onerror=&quot;alert(1)&quot;&gt;" in html
    assert "&lt;script&gt;alert(1)&lt;/script&gt;" in html
